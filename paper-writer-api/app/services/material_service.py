"""参考资料上传解析：保存文件并提取文本，供 AI 生成时参考。

支持格式（对应 aiunipaper）：
- txt / docx：提取文本
- xls / xlsx：用 openpyxl 提取单元格文本（表格/SPSS 结果）
- jpg / jpeg / png：仅保存（图片无法解析文本，作为素材留存）

单文件大小上限 5MB，总数上限 5 个。
"""

from __future__ import annotations

import logging
from pathlib import Path

from fastapi import UploadFile

logger = logging.getLogger(__name__)

ALLOWED_EXT = {"txt", "docx", "xls", "xlsx", "jpg", "jpeg", "png"}
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
MAX_FILES = 5
PER_FILE_TEXT_CHARS = 6000  # 单文件最多注入的文本字符数
TOTAL_TEXT_CHARS = 30000  # 全部资料合计最多注入的字符数
MATERIAL_KINDS = frozenset({"开题报告", "仿写论文", "其他资料"})
DEFAULT_MATERIAL_KIND = "其他资料"


def normalize_kind(value: object) -> str:
    """Return a known display category before it is ever used in a path."""
    if isinstance(value, str):
        candidate = value.strip()
        if candidate in MATERIAL_KINDS:
            return candidate
    return DEFAULT_MATERIAL_KIND


def is_supported(filename: str) -> bool:
    return Path(filename).suffix.lower().lstrip(".") in ALLOWED_EXT


def _extract_txt(path: Path) -> str:
    for enc in ("utf-8", "gb18030", "utf-16"):
        try:
            return path.read_text(encoding=enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document

        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as exc:  # noqa: BLE001
        logger.warning("docx 解析失败 %s: %s", path.name, exc)
        return ""


def _extract_excel(path: Path) -> str:
    try:
        import openpyxl

        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        lines: list[str] = []
        for ws in wb.worksheets:
            lines.append(f"# 工作表：{ws.title}")
            for row in ws.iter_rows(values_only=True):
                vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
                if vals:
                    lines.append(" | ".join(vals))
        wb.close()
        return "\n".join(lines)
    except Exception as exc:  # noqa: BLE001
        logger.warning("Excel 解析失败 %s: %s", path.name, exc)
        return ""


def extract_text(path: Path) -> str:
    """按扩展名提取文本；不支持的格式返回空串（图片等仅保存）。"""
    ext = path.suffix.lower().lstrip(".")
    if ext == "txt":
        return _extract_txt(path)
    if ext == "docx":
        return _extract_docx(path)
    if ext in ("xls", "xlsx"):
        return _extract_excel(path)
    return ""


def save_and_extract(files: list[UploadFile],
                     kinds: list[str] | None,
                     task_dir: Path) -> list[dict]:
    """保存上传的资料到 task_dir/materials/<kind>/ 并提取文本。

    files 与 kinds 按顺序一一对应；kinds 缺省时全部归为“其他资料”。
    返回 [{kind, filename, path, text}]，供写回 GenerateRequest.materials。
    """
    if not files:
        return []
    if len(files) > MAX_FILES:
        raise ValueError(f"最多上传 {MAX_FILES} 个资料文件")

    materials_dir = task_dir / "materials"
    materials_dir.mkdir(parents=True, exist_ok=True)
    materials_root = materials_dir.resolve()

    result: list[dict] = []
    total_chars = 0
    for i, up in enumerate(files):
        filename = Path(up.filename or f"material_{i}").name
        ext = Path(filename).suffix.lower().lstrip(".")
        if ext not in ALLOWED_EXT:
            raise ValueError(
                f"文件「{filename}」格式不支持，支持："
                f"txt/docx/xls/xlsx/jpg/jpeg/png")
        if up.size is not None and up.size > MAX_FILE_SIZE:
            raise ValueError(f"文件「{filename}」超过 5MB 大小限制")
        data = up.file.read()
        if len(data) > MAX_FILE_SIZE:
            raise ValueError(f"文件「{filename}」超过 5MB 大小限制")

        raw_kind = kinds[i] if kinds and i < len(kinds) else None
        kind = normalize_kind(raw_kind)
        kind_dir = (materials_root / kind).resolve()
        try:
            kind_dir.relative_to(materials_root)
        except ValueError as exc:
            raise ValueError("资料类别目录无效") from exc
        kind_dir.mkdir(parents=True, exist_ok=True)
        target = kind_dir / filename
        # 同名文件去重
        n = 1
        while target.exists():
            target = kind_dir / f"{Path(filename).stem}_{n}{Path(filename).suffix}"
            n += 1
        target.write_bytes(data)

        text = ""
        if ext in ("txt", "docx", "xls", "xlsx"):
            tmp = Path(target)
            # 文本提取按扩展名进行，图片等跳过
            text = extract_text(tmp)[:PER_FILE_TEXT_CHARS]
            total_chars += len(text)
        result.append({
            "kind": kind,
            "filename": target.name,
            "path": str(target.relative_to(task_dir)),
            "text": text,
        })

    # 合计超限时按顺序截断
    if total_chars > TOTAL_TEXT_CHARS:
        budget = TOTAL_TEXT_CHARS
        for item in result:
            if budget <= 0:
                item["text"] = ""
                continue
            if len(item["text"]) > budget:
                item["text"] = item["text"][:budget]
            budget -= len(item["text"])
    return result


def build_materials_context(materials: list[dict]) -> str:
    """把提取出的资料文本整理成注入提示词的段落。"""
    if not materials:
        return ""
    sections: list[str] = []
    for m in materials:
        if not m.get("text"):
            sections.append(
                f"- {m.get('kind', '其他资料')}：《{m.get('filename', '')}》"
                f"（已上传，图片/表格数据请人工核验）")
        else:
            sections.append(
                f"- {m.get('kind', '其他资料')}《{m.get('filename', '')}》内容：\n"
                f"{m['text']}")
    return "\n\n".join(sections)
