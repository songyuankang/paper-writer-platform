"""论文生成结果预览服务。

解析生成的 论文.docx（标题/章节/段落/表格/图片/参考文献），
转换为网页可渲染的 JSON（不直接返回 docx）。
"""

from __future__ import annotations

import html
import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

CJK_RE = re.compile(r"[\u3400-\u4dbf\u4e00-\u9fff]")
FIGURE_NUM_RE = re.compile(r"^图(\d+-\d+)\s*(.*)")


def _norm(text: str) -> str:
    return re.sub(r"[\s\u3000]+", "", text or "")


def _heading_level(p: Paragraph) -> int | None:
    name = p.style.name or ""
    m = re.search(r"(标题|Heading)\s*([1-9])", name, re.I)
    if m:
        return int(m.group(2))
    pPr = p._p.find(qn("w:pPr"))
    if pPr is not None:
        ol = pPr.find(qn("w:outlineLvl"))
        if ol is not None:
            try:
                return int(ol.get(qn("w:val"))) + 1
            except (TypeError, ValueError):
                pass
    return None


def _special(text: str) -> str | None:
    key = _norm(text)
    if key == "摘要" or key.startswith("摘要："):
        return "摘要"
    if key == "关键词" or key.startswith("关键词："):
        return "关键词"
    if key.startswith("参考文献"):
        return "参考文献"
    if key == "致谢" or key.startswith("致谢"):
        return "致谢"
    return None


def _text_heading_level(text: str) -> int | None:
    """默认模式标题没有样式，按文本编号识别：第X章 / 1 / 1.1 / 1、 等。"""
    if re.match(r"^第[一二三四五六七八九十百\d]+章", text):
        return 1
    m = re.match(r"^(\d+(?:\.\d+)*)(?:\s+|[、.．])\s*\S", text)
    if m:
        return min(3, m.group(1).count(".") + 1)
    return None


def _chart_files(task_dir: Path) -> list[str]:
    charts = task_dir / "charts"
    figures = task_dir / "figures"

    def numbered(p: Path) -> int:
        m = re.search(r"(\d+)", p.stem)
        return int(m.group(1)) if m else 0

    files: list[Path] = []
    if charts.is_dir():
        files += sorted(charts.glob("figure_*.png"), key=numbered)
        files += sorted(charts.glob("fig*.png"), key=numbered)
    if figures.is_dir():
        files += sorted(figures.glob("*.png"), key=numbered)
    return [str(f.relative_to(task_dir)).replace("\\", "/") for f in files]


def _latest_docx(task_dir: Path) -> Path:
    """优先使用最新修改版本 paper_vN.docx，否则用原始 论文.docx。"""
    base = task_dir / "论文.docx"
    best = base
    best_num = 0
    for p in task_dir.glob("paper_v*.docx"):
        m = re.search(r"paper_v(\d+)\.docx", p.name)
        if m:
            n = int(m.group(1))
            if n > best_num:
                best_num = n
                best = p
    return best


def _table_html(table: Table) -> str:
    rows = []
    for row in table.rows:
        cells = ["".join(p.text for p in cell.paragraphs) for cell in row.cells]
        rows.append("<tr>" + "".join(f"<td>{html.escape(c)}</td>" for c in cells) + "</tr>")
    return f'<table class="preview-table">{"" .join(rows)}</table>'


def _count_cjk(doc) -> int:
    total = 0
    for p in doc.paragraphs:
        total += len(CJK_RE.findall(p.text))
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                total += len(CJK_RE.findall(cell.text))
    return total


def parse_preview(task_dir: Path, request: dict | None = None) -> dict:
    """解析任务目录中的论文.docx，返回网页预览 JSON。"""
    request = request or {}
    docx_path = _latest_docx(task_dir)
    doc = Document(str(docx_path))
    chart_files = _chart_files(task_dir)

    chapters: list[dict] = []
    references: list[str] = []
    current: dict | None = None
    image_index = 0
    in_references = False
    started = False  # 跳过模板封面，直到出现摘要/关键词或正文章节
    chapter_number = 0
    paragraph_counter = 0

    def new_chapter(title: str, level: int, special: bool = False) -> dict:
        nonlocal chapter_number, paragraph_counter
        if not special and level == 1:
            chapter_number += 1
            number = str(chapter_number)
        else:
            number = ""
        paragraph_counter = 0
        return {
            "id": _slug(title) or f"sec{len(chapters) + 1}",
            "level": 0 if special else level,
            "number": number,
            "title": title,
            "content": "",
            "blocks": [],
            "images": [],
        }

    def close_chapter() -> None:
        nonlocal current
        if current is not None:
            chapters.append(current)
            current = None

    for child in doc.element.body:
        tag = child.tag.split("}")[-1]
        if tag == "tbl":
            if current is not None:
                html_table = _table_html(Table(child, doc))
                current["content"] += html_table
                current["blocks"].append(
                    {"type": "table", "html": html_table})
            continue
        if tag != "p":
            continue

        p = Paragraph(child, doc)
        text = p.text.strip()
        level = _heading_level(p) or _text_heading_level(text)
        special = _special(text)
        has_image = bool(p._p.xpath(".//w:drawing") or p._p.xpath(".//w:pict"))

        # 图片（封面前的图片也忽略）
        if has_image and not started:
            continue

        if has_image and image_index < len(chart_files):
            image = {"path": chart_files[image_index], "number": "", "title": ""}
            image_index += 1
            if current is None:
                current = new_chapter("正文", 1)
                chapters.append(current)
            idx = len(current["images"])
            current["images"].append(image)
            current["content"] += (
                f'<figure class="preview-figure" data-index="{idx}"></figure>'
            )
            current["blocks"].append({"type": "figure", **image})
            continue

        # 图题（紧跟图片后）
        if started and text and re.match(r"^图\d+-\d+", text):
            m = FIGURE_NUM_RE.match(text)
            if current is not None and current["images"]:
                current["images"][-1]["number"] = m.group(1)
                current["images"][-1]["title"] = m.group(2).strip()
            if current is not None and current["blocks"] and \
                    current["blocks"][-1].get("type") == "figure":
                current["blocks"][-1]["number"] = m.group(1)
                current["blocks"][-1]["title"] = m.group(2).strip()
            continue

        # 参考文献区
        if special == "参考文献":
            in_references = True
            started = True
            close_chapter()
            current = new_chapter("参考文献", 1, special=True)
            continue
        if in_references:
            if special == "致谢" or (level == 1 and special is None):
                in_references = False
            else:
                if text:
                    references.append(text)
                continue

        # 摘要
        if special == "摘要":
            started = True
            close_chapter()
            current = new_chapter("摘要", 0, special=True)
            continue

        # 关键词（“关键词：xxx”内容行或独立标题）
        if special == "关键词":
            content_line = bool(re.match(r"^关键词[：:]", text))
            if content_line and current is not None and current["title"] == "关键词":
                current["content"] += f"<p>{html.escape(text)}</p>"
                continue
            started = True
            close_chapter()
            current = new_chapter("关键词", 0, special=True)
            if content_line:
                current["content"] += f"<p>{html.escape(text)}</p>"
            continue

        # 标题
        if level in (1, 2, 3):
            if level == 1:
                if not started and not re.match(r"^(第.+章|\d+[.、\s])", text):
                    continue  # 封面标题，忽略
                started = True
                close_chapter()
                current = new_chapter(text, 1)
            elif current is not None:
                current["content"] += f"<h{level}>{html.escape(text)}</h{level}>"
                current["blocks"].append({
                    "id": f"ch{chapter_number}-b{len(current['blocks']) + 1}",
                    "type": f"h{level}",
                    "text": text,
                })
            continue

        # 正文段落
        if text and started and current is not None:
            current["content"] += f"<p>{html.escape(text)}</p>"
            paragraph_counter += 1
            current["blocks"].append({
                "id": f"ch{chapter_number}-p{paragraph_counter}",
                "type": "p",
                "text": text,
            })

    close_chapter()

    # 元数据
    ref_check = task_dir / "ReferenceCheck.md"
    if ref_check.exists():
        check_text = ref_check.read_text(encoding="utf-8")
        format_check = "通过" if "总体合格" in check_text else "需处理"
    else:
        format_check = "未检查"

    metadata = {
        "word_count": _count_cjk(doc),
        "target_word_count": request.get("word_count", 0),
        "chart_count": image_index,
        "reference_count": len(references),
        "format_check": format_check,
        "major": request.get("major", ""),
        "paper_type": request.get("paper_type", ""),
        "generation_mode": request.get("generation_mode", "auto"),
        "reference_style": request.get("reference_style", ""),
        "template": (task_dir.parent.parent / "uploads" /
                     f"{task_dir.name}.docx").exists(),
        "special_requirements": request.get("special_requirements") or None,
    }

    return {
        "title": request.get("title") or (docx_path.stem),
        "metadata": metadata,
        "chapters": chapters,
        "references": references,
    }


def parse_draft_preview(task_dir: Path, request: dict | None = None) -> dict:
    """预览创作向导直接生成的 draft.json（尚未导出 DOCX）。"""
    draft_path = task_dir / "draft.json"
    draft = json.loads(draft_path.read_text(encoding="utf-8"))
    request = request or {}
    chapters = []
    word_count = 0
    for section in draft.get("sections", []):
        paragraphs = [str(p.get("text") or "").strip()
                      for p in section.get("paragraphs", [])]
        paragraphs = [p for p in paragraphs if p]
        word_count += sum(len(p) for p in paragraphs)
        blocks = [{"type": "p", "text": p} for p in paragraphs]
        content = "".join(f"<p>{html.escape(p)}</p>" for p in paragraphs)
        chapters.append({
            "id": section.get("id", section.get("number", "section")),
            "level": int(section.get("level", 2)),
            "number": section.get("number", ""),
            "title": section.get("title", ""),
            "content": content,
            "blocks": blocks,
            "images": [],
        })
    refs = list(draft.get("references") or [])
    meta = draft.get("meta") or {}
    return {
        "title": draft.get("title") or request.get("title") or "论文草稿",
        "metadata": {
            "word_count": word_count,
            "target_word_count": meta.get("word_count", 0),
            "chart_count": 0,
            "reference_count": len(refs),
            "format_check": "未导出 DOCX",
            "major": meta.get("major", ""),
            "paper_type": meta.get("paper_type", ""),
            "reference_style": meta.get("reference_style", ""),
            "template": False,
            "special_requirements": meta.get("special_requirements"),
        },
        "chapters": chapters,
        "references": refs,
    }


def _slug(text: str) -> str:
    slug = re.sub(r"[^\w\u4e00-\u9fff]+", "-", text).strip("-")
    return slug[:40]


def load_request(task_dir: Path) -> dict:
    path = task_dir / "request.json"
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:  # noqa: BLE001
        return {}
