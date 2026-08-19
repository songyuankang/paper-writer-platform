"""格式处理任务：创建/执行/状态/下载。

输出：outputs/<task_id>/formatted/paper.docx、paper.pdf（可选）、format_report.md。
"""

from __future__ import annotations

import json
import logging
import shutil
import subprocess
import threading
import uuid
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.config import settings
from app.db import get_conn
from app.formatter import docx_builder, reference as reference_mod
from app.formatter import template_manager
from app.formatter.service import spec_from_paper_content

logger = logging.getLogger(__name__)


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


class FormatTaskManager:
    def create(self, task_id: str, template_id: str | None,
               settings_data: dict) -> str:
        format_id = uuid.uuid4().hex
        with get_conn() as conn:
            conn.execute(
                """
                INSERT INTO format_tasks
                    (id, task_id, template_id, status, progress, message,
                     settings, created_at)
                VALUES (?, ?, ?, 'waiting', 0, ?, ?, ?)
                """,
                (format_id, task_id, template_id, "等待处理",
                 json.dumps(settings_data or {}, ensure_ascii=False), _now()),
            )
        return format_id

    def start(self, format_id: str) -> bool:
        if self.get(format_id) is None:
            return False
        threading.Thread(target=self.run, args=(format_id,),
                         daemon=True).start()
        return True

    def get(self, format_id: str) -> dict | None:
        with get_conn() as conn:
            row = conn.execute(
                "SELECT * FROM format_tasks WHERE id = ?", (format_id,)).fetchone()
        if row is None:
            return None
        record = dict(row)
        if record.get("settings"):
            try:
                record["settings"] = json.loads(record["settings"])
            except json.JSONDecodeError:
                record["settings"] = {}
        if record.get("files"):
            try:
                record["files"] = json.loads(record["files"])
            except json.JSONDecodeError:
                record["files"] = []
        return record

    def _update(self, format_id: str, status: str, progress: int,
                message: str, files: list[str] | None = None) -> None:
        with get_conn() as conn:
            conn.execute(
                "UPDATE format_tasks SET status=?, progress=?, message=?, "
                "files=?, completed_at=COALESCE(completed_at, ?) WHERE id=?",
                (status, progress, message,
                 json.dumps(files or [], ensure_ascii=False),
                 _now() if status in ("completed", "failed") else None,
                 format_id),
            )

    def run(self, format_id: str) -> None:
        record = self.get(format_id)
        if record is None:
            return
        task_id = record["task_id"]
        task_dir = settings.output_dir / task_id
        try:
            self._update(format_id, "processing", 10, "读取论文内容")
            paper_info = _load_request(task_dir)
            content_dir = task_dir / "paper_content"
            if content_dir.is_dir():
                spec = spec_from_paper_content(paper_info, content_dir)
            else:
                spec = json.loads(
                    (task_dir / "paper_spec.json").read_text(encoding="utf-8"))

            settings_data = record.get("settings") or {}
            ref_style = settings_data.get("reference_style") or \
                paper_info.get("reference_style", "gb7714")
            use_template = bool(settings_data.get("use_template", True))
            template_path = None
            if use_template and record.get("template_id"):
                template_path = template_manager.style_path(record["template_id"])
                if template_path is None:
                    # New v2 uploads keep their original DOCX next to the
                    # structured template.  The legacy formatter can use the
                    # same source file during the migration period.
                    from app.formatter.template import get_service
                    template_path = get_service().repo.source_docx_path(
                        record["template_id"])

            self._update(format_id, "processing", 40, "正在排版 Word 文档")
            out_dir = task_dir / "formatted"
            out_dir.mkdir(parents=True, exist_ok=True)
            docx_name = "formatted/paper.docx"
            docx_builder.build_docx(
                task_dir, spec, spec.get("meta", {}), docx_name,
                template_path=template_path,
                toc_update=bool(settings_data.get("toc", True)))
            # 参考文献格式化与检查
            reference_mod.write_reference_deliverables(task_dir, spec, ref_style)

            self._update(format_id, "checking", 70, "正在检查格式")
            docx_path = task_dir / docx_name
            pdf_note = _try_pdf(docx_path)
            report = _build_format_report(
                task_dir, docx_path, spec,
                template_path is not None, settings_data, pdf_note)
            (out_dir / "format_report.md").write_text(report, encoding="utf-8")

            files = [f"formatted/{p.name}" for p in sorted(out_dir.iterdir())
                     if p.is_file()]
            self._update(format_id, "completed", 100, "格式处理完成", files)
        except Exception as exc:  # noqa: BLE001
            logger.exception("Format task %s failed", format_id)
            self._update(format_id, "failed", 0, f"{type(exc).__name__}: {exc}")

    def formatted_dir(self, format_id: str) -> Path | None:
        record = self.get(format_id)
        if record is None:
            return None
        return settings.output_dir / record["task_id"] / "formatted"


def _load_request(task_dir: Path) -> dict:
    path = task_dir / "request.json"
    if path.exists():
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            return {}
    return {}


def _try_pdf(docx_path: Path) -> str:
    exe = shutil.which("soffice") or shutil.which("libreoffice")
    if not exe:
        return "PDF 未生成（未检测到 LibreOffice/soffice，可在环境安装后重新格式化）"
    try:
        subprocess.run(
            [exe, "--headless", "--convert-to", "pdf",
             "--outdir", str(docx_path.parent), str(docx_path)],
            timeout=180, capture_output=True)
        pdf = docx_path.with_suffix(".pdf")
        if pdf.exists():
            return "PDF 已生成"
    except Exception as exc:  # noqa: BLE001
        return f"PDF 生成失败：{exc}"
    return "PDF 未生成"


def _build_format_report(task_dir: Path, docx_path: Path, spec: dict,
                         used_template: bool, settings_data: dict,
                         pdf_note: str) -> str:
    doc = Document(str(docx_path))
    toc_found = any(
        p._p.xpath(".//w:fldSimple[contains(@w:instr,'TOC')]") or
        p._p.xpath(".//w:instrText[contains(.,'TOC')]")
        for p in doc.paragraphs)
    has_east_asia = any(
        (r._element.rPr is not None and r._element.rPr.rFonts is not None and
         r._element.rPr.rFonts.get(qn("w:eastAsia")))
        for p in doc.paragraphs for r in p.runs)
    refs = spec.get("references", [])
    figures = len(doc.inline_shapes)
    chart_dir = task_dir / "charts"
    charts = figures > 0 or (chart_dir.is_dir() and list(chart_dir.glob("*.png")))

    def status(ok: bool, extra: str = "") -> str:
        return f"{'✓ 通过' if ok else '✗ 失败'} {extra}".rstrip()

    lines = [
        "# 格式处理报告（FormatReport）",
        "",
        f"- 模板：{'学校模板' if used_template else '默认模板'}",
        f"- 目录设置：{'自动生成' if settings_data.get('toc', True) else '关闭'}",
        f"- 参考文献格式：{settings_data.get('reference_style', 'gb7714')}",
        "",
        "## 检查结果",
        "",
        "| 项目 | 结果 |",
        "| --- | --- |",
        f"| 字体 | {status(has_east_asia, '（检测到中文字体设置）')} |",
        f"| 目录 | {status(toc_found, '（已生成目录域）' if toc_found else '（未检测到目录域）')} |",
        f"| 参考文献 | {status(bool(refs), f'（{len(refs)} 条）')} |",
        f"| 图表 | {status(bool(charts), f'（{figures} 张内嵌图）')} |",
        "",
        "## 说明",
        "",
        f"- {pdf_note}",
        "- 如需 PDF 请在部署环境安装 LibreOffice 后重新执行格式处理",
    ]
    return "\n".join(lines) + "\n"


format_manager = FormatTaskManager()
