"""模板渲染器（TemplateRenderer）：Template → DOCX。

把 v2 ``Template``（meta/page/header/footer/numbering/blocks）与论文内容
（沿用现有 ``paper_spec`` 结构：sections 的 h1-h4 / p / table / figure /
pagebreak / references）渲染为 python-docx ``Document``。

设计原则：
- **样式唯一入口**：所有 ``TemplateStyle`` → DOCX 的应用集中在
  :meth:`TemplateRenderer._apply_style`，各渲染方法不重新硬编码字体/字号；
- **首行缩进**：``first_line_indent.unit == "chars"`` 时写 Word OOXML
  ``w:firstLineChars``（精确按字符数），不用 Pt 估算；
- **标题编号**：``numbering.enabled`` 时按模板 ``h1..h4`` 模式生成编号；
  关闭时保持 spec 原始标题；已带编号（"1 引言"/"第一章"）不重复编号；
- **缺失 block 兜底**：当前模板 → 默认模板（``default_template``）→
  ``SAFE_DEFAULTS``，任何 block 缺失不导致渲染失败；
- **TOC**：真实 Word TOC 域（``w:fldSimple``/``w:instrText TOC``）+ 设置
  ``w:updateFields`` 打开时更新；heading 正确设置 ``w:outlineLvl``；
- **参考文献内容格式化**：复用现有 ``reference.py``（GB/T 7714 等）逻辑；
- **cover.docx**：本阶段不实现（无 cover.docx 的完整渲染），预留
  :meth:`_open_document` 扩展接口，不为此改动现有架构；
- 本模块独立，不修改/不依赖旧 docx_builder / template_manager / style 流程。

约定：``_apply_style`` 需要同时拿到 paragraph 与 run（段落属性与字符属性）。
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from app.formatter.template.models import (
    DEFAULT_EAST_ASIA_FONT,
    DEFAULT_LATIN_FONT,
    IndentUnit,
    LineSpacingMode,
    Template,
    TemplateBlock,
    TemplateStyle,
    TextAlign,
)

#: 纸张规格：名称 → (宽 mm, 高 mm)
PAGE_SIZES = {
    "A3": (297, 420),
    "A4": (210, 297),
    "A5": (148, 210),
    "B4": (250, 353),
    "B5": (176, 250),
    "Letter": (216, 279),
    "Legal": (216, 356),
    "Tabloid": (279, 432),
}

#: 自定义纸张尺寸 "210x297mm"
_CUSTOM_SIZE_RE = re.compile(r"^(\d+(?:\.\d+)?)x(\d+(?:\.\d+)?)mm$", re.I)

#: 默认编号模式（与 v2 内置模板一致）
_DEFAULT_NUMBERING_PATTERNS = {
    1: "第{chinese}章",
    2: "{h1}.{n}",
    3: "{h1}.{h2}.{n}",
    4: "{h1}.{h2}.{h3}.{n}",
}

#: 已带编号前缀（"1 引言" / "1.1 研究" / "第一章" / "1、"）
_NUMBERED_PREFIX = re.compile(
    r"^\s*(?:第[零一二三四五六七八九十百\d]+[章篇卷]"
    r"|\d+(?:[\.．]\d+)*[\s、，,．.])")

#: 特殊一级标题 → 对应 block（不参与自动编号）
_SPECIAL_HEADINGS = {
    "致谢": ("acknowledgement", "title"),
    "参考文献": ("references", "title"),
    "附录": ("appendix", "title"),
}

#: 中文数字
_CN_DIGITS = "零一二三四五六七八九"


def _cn_num(n: int) -> str:
    """阿拉伯数字 → 中文数字（1-99 精确，更大回退阿拉伯）。"""
    if n <= 0:
        return str(n)
    if n < 10:
        return _CN_DIGITS[n]
    if n < 20:
        return "十" + (_CN_DIGITS[n - 10] if n > 10 else "")
    if n < 100:
        tens, ones = divmod(n, 10)
        return _CN_DIGITS[tens] + "十" + (_CN_DIGITS[ones] if ones else "")
    return str(n)


#: 缺失 block 的安全兜底（与 v2 内置模板风格一致的最小集合）
SAFE_DEFAULTS: dict[str, dict] = {
    "title_zh": {
        "key": "title_zh", "kind": "title_zh", "label": "论文标题", "enabled": True,
        "styles": {"self": {
            "font_family": {"east_asia": "黑体", "latin": "Times New Roman"},
            "font_size_pt": 22, "bold": True, "italic": False, "underline": False,
            "alignment": "center", "line_spacing": {"mode": "multiple", "value": 1.25},
            "space_before_pt": 6, "space_after_pt": 6,
            "first_line_indent": {"unit": "chars", "value": 0},
            "keep_with_next": True, "page_break_before": False}},
    },
    "title_en": {
        "key": "title_en", "kind": "title_en", "label": "英文标题", "enabled": True,
        "styles": {"self": {
            "font_family": {"east_asia": "Times New Roman", "latin": "Times New Roman"},
            "font_size_pt": 16, "bold": True, "italic": False, "underline": False,
            "alignment": "center", "line_spacing": {"mode": "multiple", "value": 1.25},
            "space_before_pt": 6, "space_after_pt": 6,
            "first_line_indent": {"unit": "chars", "value": 0},
            "keep_with_next": True, "page_break_before": False}},
    },
    "abstract": {
        "key": "abstract", "kind": "abstract", "label": "摘要", "enabled": True,
        "styles": {
            "title": {
                "font_family": {"east_asia": "黑体", "latin": "Times New Roman"},
                "font_size_pt": 15, "bold": True, "italic": False, "underline": False,
                "alignment": "center", "line_spacing": {"mode": "multiple", "value": 1.25},
                "space_before_pt": 6, "space_after_pt": 6,
                "first_line_indent": {"unit": "chars", "value": 0},
                "keep_with_next": True, "page_break_before": False},
            "content": {
                "font_family": {"east_asia": "宋体", "latin": "Times New Roman"},
                "font_size_pt": 12, "bold": False, "italic": False, "underline": False,
                "alignment": "justify", "line_spacing": {"mode": "multiple", "value": 1.5},
                "space_before_pt": 0, "space_after_pt": 0,
                "first_line_indent": {"unit": "chars", "value": 2},
                "keep_with_next": False, "page_break_before": False}},
    },
    "keywords": {
        "key": "keywords", "kind": "keywords", "label": "关键词", "enabled": True,
        "styles": {
            "label": {
                "font_family": {"east_asia": "黑体", "latin": "Times New Roman"},
                "font_size_pt": 12, "bold": True, "italic": False, "underline": False,
                "alignment": "justify", "line_spacing": {"mode": "multiple", "value": 1.5},
                "space_before_pt": 6, "space_after_pt": 0,
                "first_line_indent": {"unit": "chars", "value": 0},
                "keep_with_next": False, "page_break_before": False},
            "content": {
                "font_family": {"east_asia": "宋体", "latin": "Times New Roman"},
                "font_size_pt": 12, "bold": False, "italic": False, "underline": False,
                "alignment": "justify", "line_spacing": {"mode": "multiple", "value": 1.5},
                "space_before_pt": 0, "space_after_pt": 0,
                "first_line_indent": {"unit": "chars", "value": 0},
                "keep_with_next": False, "page_break_before": False}},
    },
    "heading1": {
        "key": "heading1", "kind": "heading", "label": "一级标题", "level": 1,
        "enabled": True, "styles": {"self": {
            "font_family": {"east_asia": "黑体", "latin": "Times New Roman"},
            "font_size_pt": 16, "bold": True, "italic": False, "underline": False,
            "alignment": "left", "line_spacing": {"mode": "multiple", "value": 1.25},
            "space_before_pt": 12, "space_after_pt": 12,
            "first_line_indent": {"unit": "chars", "value": 0},
            "keep_with_next": True, "page_break_before": False}},
    },
    "heading2": {
        "key": "heading2", "kind": "heading", "label": "二级标题", "level": 2,
        "enabled": True, "styles": {"self": {
            "font_family": {"east_asia": "黑体", "latin": "Times New Roman"},
            "font_size_pt": 14, "bold": True, "italic": False, "underline": False,
            "alignment": "left", "line_spacing": {"mode": "multiple", "value": 1.25},
            "space_before_pt": 6, "space_after_pt": 6,
            "first_line_indent": {"unit": "chars", "value": 0},
            "keep_with_next": True, "page_break_before": False}},
    },
    "heading3": {
        "key": "heading3", "kind": "heading", "label": "三级标题", "level": 3,
        "enabled": True, "styles": {"self": {
            "font_family": {"east_asia": "黑体", "latin": "Times New Roman"},
            "font_size_pt": 12, "bold": True, "italic": False, "underline": False,
            "alignment": "left", "line_spacing": {"mode": "multiple", "value": 1.25},
            "space_before_pt": 6, "space_after_pt": 6,
            "first_line_indent": {"unit": "chars", "value": 0},
            "keep_with_next": True, "page_break_before": False}},
    },
    "heading4": {
        "key": "heading4", "kind": "heading", "label": "四级标题", "level": 4,
        "enabled": True, "styles": {"self": {
            "font_family": {"east_asia": "黑体", "latin": "Times New Roman"},
            "font_size_pt": 12, "bold": True, "italic": False, "underline": False,
            "alignment": "left", "line_spacing": {"mode": "multiple", "value": 1.25},
            "space_before_pt": 6, "space_after_pt": 6,
            "first_line_indent": {"unit": "chars", "value": 0},
            "keep_with_next": True, "page_break_before": False}},
    },
    "paragraph": {
        "key": "body", "kind": "paragraph", "label": "正文", "enabled": True,
        "styles": {"self": {
            "font_family": {"east_asia": "宋体", "latin": "Times New Roman"},
            "font_size_pt": 12, "bold": False, "italic": False, "underline": False,
            "alignment": "justify", "line_spacing": {"mode": "multiple", "value": 1.5},
            "space_before_pt": 0, "space_after_pt": 0,
            "first_line_indent": {"unit": "chars", "value": 2},
            "keep_with_next": False, "page_break_before": False}},
    },
    "figure_caption": {
        "key": "figure_caption", "kind": "figure_caption", "label": "图题注",
        "enabled": True, "styles": {"self": {
            "font_family": {"east_asia": "黑体", "latin": "Times New Roman"},
            "font_size_pt": 10.5, "bold": True, "italic": False, "underline": False,
            "alignment": "center", "line_spacing": {"mode": "multiple", "value": 1.25},
            "space_before_pt": 3, "space_after_pt": 6,
            "first_line_indent": {"unit": "chars", "value": 0},
            "keep_with_next": False, "page_break_before": False}},
    },
    "table_caption": {
        "key": "table_caption", "kind": "table_caption", "label": "表题注",
        "enabled": True, "styles": {"self": {
            "font_family": {"east_asia": "黑体", "latin": "Times New Roman"},
            "font_size_pt": 10.5, "bold": True, "italic": False, "underline": False,
            "alignment": "center", "line_spacing": {"mode": "multiple", "value": 1.25},
            "space_before_pt": 6, "space_after_pt": 3,
            "first_line_indent": {"unit": "chars", "value": 0},
            "keep_with_next": False, "page_break_before": False}},
    },
    "toc": {
        "key": "toc", "kind": "toc", "label": "目录", "enabled": True,
        "styles": {"title": {
            "font_family": {"east_asia": "黑体", "latin": "Times New Roman"},
            "font_size_pt": 16, "bold": True, "italic": False, "underline": False,
            "alignment": "center", "line_spacing": {"mode": "multiple", "value": 1.5},
            "space_before_pt": 24, "space_after_pt": 24,
            "first_line_indent": {"unit": "chars", "value": 0},
            "keep_with_next": True, "page_break_before": False}},
    },
    "references": {
        "key": "references", "kind": "references", "label": "参考文献",
        "enabled": True, "settings": {"style": "gb7714"},
        "styles": {
            "title": {
                "font_family": {"east_asia": "黑体", "latin": "Times New Roman"},
                "font_size_pt": 16, "bold": True, "italic": False, "underline": False,
                "alignment": "center", "line_spacing": {"mode": "multiple", "value": 1.5},
                "space_before_pt": 24, "space_after_pt": 18,
                "first_line_indent": {"unit": "chars", "value": 0},
                "keep_with_next": True, "page_break_before": False},
            "item": {
                "font_family": {"east_asia": "宋体", "latin": "Times New Roman"},
                "font_size_pt": 10.5, "bold": False, "italic": False,
                "underline": False, "alignment": "left",
                "line_spacing": {"mode": "multiple", "value": 1.25},
                "space_before_pt": 3, "space_after_pt": 0,
                "first_line_indent": {"unit": "chars", "value": 0},
                "keep_with_next": False, "page_break_before": False}},
    },
    "acknowledgement": {
        "key": "acknowledgement", "kind": "acknowledgement", "label": "致谢",
        "enabled": True,
        "styles": {
            "title": {
                "font_family": {"east_asia": "黑体", "latin": "Times New Roman"},
                "font_size_pt": 16, "bold": True, "italic": False, "underline": False,
                "alignment": "center", "line_spacing": {"mode": "multiple", "value": 1.5},
                "space_before_pt": 24, "space_after_pt": 18,
                "first_line_indent": {"unit": "chars", "value": 0},
                "keep_with_next": True, "page_break_before": False},
            "content": {
                "font_family": {"east_asia": "宋体", "latin": "Times New Roman"},
                "font_size_pt": 12, "bold": False, "italic": False,
                "underline": False, "alignment": "justify",
                "line_spacing": {"mode": "multiple", "value": 1.5},
                "space_before_pt": 0, "space_after_pt": 0,
                "first_line_indent": {"unit": "chars", "value": 2},
                "keep_with_next": False, "page_break_before": False}},
    },
    "appendix": {
        "key": "appendix", "kind": "appendix", "label": "附录", "enabled": True,
        "styles": {
            "title": {
                "font_family": {"east_asia": "黑体", "latin": "Times New Roman"},
                "font_size_pt": 16, "bold": True, "italic": False, "underline": False,
                "alignment": "center", "line_spacing": {"mode": "multiple", "value": 1.5},
                "space_before_pt": 24, "space_after_pt": 18,
                "first_line_indent": {"unit": "chars", "value": 0},
                "keep_with_next": True, "page_break_before": False},
            "content": {
                "font_family": {"east_asia": "宋体", "latin": "Times New Roman"},
                "font_size_pt": 12, "bold": False, "italic": False,
                "underline": False, "alignment": "justify",
                "line_spacing": {"mode": "multiple", "value": 1.5},
                "space_before_pt": 0, "space_after_pt": 0,
                "first_line_indent": {"unit": "chars", "value": 2},
                "keep_with_next": False, "page_break_before": False}},
    },
    "abstract_en": {
        "key": "abstract_en", "kind": "abstract_en", "label": "Abstract",
        "enabled": True,
        "styles": {
            "title": {
                "font_family": {"east_asia": "Times New Roman",
                                "latin": "Times New Roman"},
                "font_size_pt": 15, "bold": True, "italic": False,
                "underline": False, "alignment": "center",
                "line_spacing": {"mode": "multiple", "value": 1.25},
                "space_before_pt": 6, "space_after_pt": 6,
                "first_line_indent": {"unit": "chars", "value": 0},
                "keep_with_next": True, "page_break_before": False},
            "content": {
                "font_family": {"east_asia": "Times New Roman",
                                "latin": "Times New Roman"},
                "font_size_pt": 12, "bold": False, "italic": False,
                "underline": False, "alignment": "justify",
                "line_spacing": {"mode": "multiple", "value": 1.5},
                "space_before_pt": 0, "space_after_pt": 0,
                "first_line_indent": {"unit": "chars", "value": 0},
                "keep_with_next": False, "page_break_before": False}},
    },
    "keywords_en": {
        "key": "keywords_en", "kind": "keywords_en", "label": "Keywords",
        "enabled": True,
        "styles": {
            "label": {
                "font_family": {"east_asia": "Times New Roman",
                                "latin": "Times New Roman"},
                "font_size_pt": 12, "bold": True, "italic": False,
                "underline": False, "alignment": "justify",
                "line_spacing": {"mode": "multiple", "value": 1.5},
                "space_before_pt": 6, "space_after_pt": 0,
                "first_line_indent": {"unit": "chars", "value": 0},
                "keep_with_next": False, "page_break_before": False},
            "content": {
                "font_family": {"east_asia": "Times New Roman",
                                "latin": "Times New Roman"},
                "font_size_pt": 12, "bold": False, "italic": False,
                "underline": False, "alignment": "justify",
                "line_spacing": {"mode": "multiple", "value": 1.5},
                "space_before_pt": 0, "space_after_pt": 0,
                "first_line_indent": {"unit": "chars", "value": 0},
                "keep_with_next": False, "page_break_before": False}},
    },
}


class TemplateRenderer:
    """Template → DOCX 渲染器（无状态，可复用）。

    ``default_template`` 用于缺失 block 的第二级兜底（当前模板 → 默认模板 →
    SAFE_DEFAULTS）。不传时跳过默认模板，直接回退 SAFE_DEFAULTS。
    """

    def __init__(self, default_template: Template | None = None):
        self.default_template = default_template
        self._counters: list[int] = [0, 0, 0, 0]
        self._source_heading_styles: tuple[str, ...] = ()

    # ------------------------------------------------------------------
    # 入口
    # ------------------------------------------------------------------
    def render(self, template: Template, spec: dict,
               base_dir: Path | str | None = None,
               source_docx: Path | str | None = None) -> Document:
        """渲染模板 + 内容 → 内存 Document（不落盘）。"""
        self._counters = [0, 0, 0, 0]
        self._numbering = template.numbering or {}
        self._source_heading_styles = self._source_heading_style_names(source_docx)
        doc = self._open_document(source_docx)
        self._apply_page_setup(doc, template.page)
        self._apply_header_footer(doc, template.header, template.footer)
        self._render_content(doc, spec, template, base_dir)
        return doc

    def render_document(self, template: Template, spec: dict,
                        task_dir: Path | str,
                        out_name: str = "论文.docx") -> Path:
        """兼容入口：渲染并保存 docx，返回输出路径。

        内部 = :meth:`render`（Template → Document）→ ``DocxExporter``
        （Document → .docx 文件）。文件名安全处理 / 目录创建 / 覆盖策略 /
        ``document.save`` 全部交给 DocxExporter，Renderer 不直接落盘。
        """
        from app.formatter.template.exporter import DocxExporter
        doc = self.render(template, spec, base_dir=task_dir)
        return DocxExporter().export(doc, task_dir, out_name)

    # ------------------------------------------------------------------
    # 文档打开（cover.docx 扩展点）
    # ------------------------------------------------------------------
    def _open_document(self, source_docx: Path | str | None = None) -> Document:
        """新建空白文档。

        未来支持 cover.docx 时：若模板目录存在 cover.docx，则以它为基底
        打开（保留封面/页眉页脚母版），并定位正文插入点；本阶段先新建。
        """
        path = Path(source_docx) if source_docx else None
        if path is None or not path.is_file():
            return Document()
        doc = Document(str(path))
        body = doc.element.body
        sect_pr = body.sectPr
        for child in list(body):
            if child is not sect_pr:
                body.remove(child)
        return doc

    # ------------------------------------------------------------------
    # 页面设置
    # ------------------------------------------------------------------
    def _apply_page_setup(self, doc: Document, page: dict) -> None:
        if not isinstance(page, dict) or not page:
            return
        sec = doc.sections[0]
        size = page.get("size") or "A4"
        w_mm, h_mm = self._page_size_mm(size)
        orientation = page.get("orientation", "portrait")
        if orientation == "landscape" and w_mm < h_mm:
            w_mm, h_mm = h_mm, w_mm
        sec.orientation = (WD_ORIENT.LANDSCAPE if w_mm > h_mm
                           else WD_ORIENT.PORTRAIT)
        sec.page_width = Mm(w_mm)
        sec.page_height = Mm(h_mm)

        margins = page.get("margins")
        if isinstance(margins, dict):
            if _is_num(margins.get("top_mm")):
                sec.top_margin = Mm(margins["top_mm"])
            if _is_num(margins.get("bottom_mm")):
                sec.bottom_margin = Mm(margins["bottom_mm"])
            if _is_num(margins.get("left_mm")):
                sec.left_margin = Mm(margins["left_mm"])
            if _is_num(margins.get("right_mm")):
                sec.right_margin = Mm(margins["right_mm"])
        if _is_num(page.get("header_distance_mm")):
            sec.header_distance = Mm(page["header_distance_mm"])
        if _is_num(page.get("footer_distance_mm")):
            sec.footer_distance = Mm(page["footer_distance_mm"])

    @staticmethod
    def _page_size_mm(size: str) -> tuple[float, float]:
        if size in PAGE_SIZES:
            return PAGE_SIZES[size]
        m = _CUSTOM_SIZE_RE.match(size or "")
        if m:
            return float(m.group(1)), float(m.group(2))
        return PAGE_SIZES["A4"]

    # ------------------------------------------------------------------
    # 页眉 / 页脚（含 {page} → Word PAGE 域）
    # ------------------------------------------------------------------
    def _apply_header_footer(self, doc: Document, header: dict,
                             footer: dict) -> None:
        sec = doc.sections[0]
        if isinstance(header, dict) and header.get("content"):
            paragraph = sec.header.paragraphs[0]
            paragraph.clear()
            self._render_field_text(
                paragraph, str(header["content"]),
                TemplateStyle.from_dict(header.get("style")))
        if isinstance(footer, dict) and footer.get("content"):
            paragraph = sec.footer.paragraphs[0]
            paragraph.clear()
            self._render_field_text(
                paragraph, str(footer["content"]),
                TemplateStyle.from_dict(footer.get("style")))

    def _render_field_text(self, paragraph, content: str,
                           style: TemplateStyle) -> None:
        """把含 ``{page}`` 的内容渲染为文本 + PAGE 域。"""
        parts = re.split(r"\{page\}", content)
        for i, part in enumerate(parts):
            if part:
                run = paragraph.add_run(part)
                self._apply_style(paragraph, run, style)
            if i < len(parts) - 1:
                run = paragraph.add_run()
                self._apply_style(paragraph, run, style)
                self._insert_field(run, " PAGE ")

    @staticmethod
    def _insert_field(run, instr: str) -> None:
        """在 run 中插入 Word 域（begin + instrText + end）。"""
        r = run._r
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr_el = OxmlElement("w:instrText")
        instr_el.set(qn("xml:space"), "preserve")
        instr_el.text = instr
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r.append(fld_begin)
        r.append(instr_el)
        r.append(fld_end)

    # ------------------------------------------------------------------
    # 内容渲染（spec 分发）
    # ------------------------------------------------------------------
    def _render_content(self, doc: Document, spec: dict, template: Template,
                        base_dir: Path | None = None) -> None:
        meta = spec.get("meta") or {}
        if meta.get("title"):
            self._render_title(doc, meta["title"], template, "title_zh")
        if meta.get("title_en"):
            self._render_title(doc, meta["title_en"], template, "title_en")
        if meta.get("abstract"):
            self._render_abstract(doc, meta["abstract"], template, "abstract")
        if meta.get("keywords"):
            self._render_keywords(doc, meta["keywords"], template, "keywords")
        if meta.get("abstract_en"):
            if meta.get("abstract") or meta.get("keywords"):
                doc.add_page_break()
            self._render_abstract(doc, meta["abstract_en"], template, "abstract_en")
        if meta.get("keywords_en"):
            self._render_keywords(doc, meta["keywords_en"], template, "keywords_en")

        toc_rendered = False
        for item in spec.get("sections") or []:
            if not isinstance(item, dict):
                continue
            kind = item.get("type")
            if kind in ("h1", "h2", "h3", "h4") and not toc_rendered:
                has_prelim = bool(meta.get("abstract") or meta.get("keywords") or
                                  meta.get("abstract_en") or meta.get("keywords_en"))
                if has_prelim:
                    doc.add_page_break()
                if self._toc_enabled(template) and bool(meta.get("toc", True)):
                    self._render_toc(doc, template)
                    doc.add_page_break()
                toc_rendered = True
            if kind == "h1":
                self._render_heading(doc, item.get("text", ""), 1, template)
            elif kind == "h2":
                self._render_heading(doc, item.get("text", ""), 2, template)
            elif kind == "h3":
                self._render_heading(doc, item.get("text", ""), 3, template)
            elif kind == "h4":
                self._render_heading(doc, item.get("text", ""), 4, template)
            elif kind == "p":
                self._render_paragraph(doc, item.get("text", ""), template)
            elif kind == "figure":
                self._render_figure(doc, item, template, base_dir)
            elif kind == "table":
                self._render_table(doc, item, template)
            elif kind == "pagebreak":
                doc.add_page_break()
            elif kind in ("sectionbreak", "section_break"):
                self._render_section_break(doc, item, template)
            elif kind == "references":
                self._render_references(doc, item, template, meta)
            elif kind == "acknowledgement":
                self._render_special(doc, item.get("text") or item.get("content") or "",
                                     template, "acknowledgement")
            elif kind == "appendix":
                self._render_special(doc, item.get("text") or item.get("content") or "",
                                     template, "appendix")

    def _render_title(self, doc: Document, text: str, template: Template,
                      key: str) -> None:
        block = self._resolve_block(template, key, key)
        style = self._block_style(block, "self")
        self._add_styled_paragraph(doc, str(text), style)

    def _render_abstract(self, doc: Document, text: str, template: Template,
                         key: str) -> None:
        block = self._resolve_block(template, key, key)
        title_style = self._block_style(block, "title")
        title_style.alignment = TextAlign.CENTER
        title_style.first_line_indent_value = 0.0
        content_style = self._block_style(block, "content")
        title = {"abstract": "摘  要", "abstract_en": "Abstract"}.get(key, "摘  要")
        self._add_styled_paragraph(doc, title, title_style)
        for para in str(text).splitlines() or [str(text)]:
            if para.strip():
                self._add_styled_paragraph(doc, para, content_style)

    def _render_keywords(self, doc: Document, keywords: list | str,
                         template: Template, key: str) -> None:
        block = self._resolve_block(template, key, key)
        label_style = self._block_style(block, "label")
        content_style = self._block_style(block, "content")
        label = {"keywords": "关键词：", "keywords_en": "Keywords: "}.get(
            key, "关键词：")
        items = keywords if isinstance(keywords, list) else [keywords]
        p = doc.add_paragraph()
        run_label = p.add_run(label)
        self._apply_style(p, run_label, label_style)
        run_body = p.add_run("；".join(str(x) for x in items))
        self._apply_style(p, run_body, content_style)

    def _render_heading(self, doc: Document, text: str, level: int,
                        template: Template) -> None:
        stripped = self._normalize_heading_text((text or "").strip())
        # 特殊一级标题（致谢/参考文献/附录）→ 对应 block，不自动编号
        if level == 1 and stripped in _SPECIAL_HEADINGS:
            key, role = _SPECIAL_HEADINGS[stripped]
            block = self._resolve_block(template, key, None)
            style = self._block_style(block, role)
            self._add_styled_paragraph(doc, stripped, style,
                                       outline_level=1)
            return
        block = self._resolve_block(template, f"heading{level}", "heading")
        style = self._block_style(block, "self")
        display = stripped
        # 防止 spec 已带编号（"1 引言"/"第一章"）时重复编号
        if not self._looks_numbered(stripped):
            number = self._heading_number(level)
            if number:
                display = f"{number} {stripped}"
        self._add_styled_paragraph(doc, display, style, outline_level=level)

    def _render_paragraph(self, doc: Document, text: str,
                          template: Template) -> None:
        block = self._resolve_block(template, "body", "paragraph")
        style = self._block_style(block, "self")
        self._add_styled_paragraph(doc, str(text), style)

    def _render_toc(self, doc: Document, template: Template) -> None:
        block = self._resolve_block(template, "toc", "toc")
        title_style = self._block_style(block, "title")
        # 目录标题不能设置 outline_level —— 否则刷新后的目录会把
        # “目 录”自己收进 TOC（自引用条目）
        self._add_styled_paragraph(doc, "目  录", title_style)
        # 真正的 Word TOC 域
        p = doc.add_paragraph()
        run = p.add_run()
        r = run._r
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = ' TOC \\o "1-3" \\h \\z \\u '
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = "（打开文档后更新域以生成目录）"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r.append(fld_begin)
        r.append(instr)
        r.append(fld_sep)
        r.append(placeholder)
        r.append(fld_end)
        self._mark_toc_update(doc)

    def _render_figure(self, doc: Document, item: dict, template: Template,
                       base_dir) -> None:
        path = Path(item.get("path", ""))
        if not path.is_absolute() and base_dir is not None:
            path = Path(base_dir) / path
        if not path.is_file():
            # 图片缺失不中断
            self._render_paragraph(doc, f"[图片缺失: {item.get('path', '')}]",
                                   template)
            return
        width_mm = min(self._content_width_mm(template), 140.0)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Mm(width_mm))
        title = item.get("title")
        if title:
            block = self._resolve_block(template, "figure_caption",
                                        "figure_caption")
            style = self._block_style(block, "self")
            self._add_styled_paragraph(doc, str(title), style)

    def _render_section_break(self, doc: Document, item: dict,
                              template: Template) -> None:
        """Insert a document section using an explicit content-spec boundary."""
        block = self._resolve_block(template, item.get("template_key"), "sectionbreak")
        settings = dict(block.settings or {})
        settings.update({k: v for k, v in item.items() if k in {"start_type", "page", "header", "footer", "page_number"}})
        start_type = str(settings.get("start_type") or "NEW_PAGE").upper()
        section_type = {
            "CONTINUOUS": WD_SECTION.CONTINUOUS,
            "EVEN_PAGE": WD_SECTION.EVEN_PAGE,
            "ODD_PAGE": WD_SECTION.ODD_PAGE,
            "NEW_PAGE": WD_SECTION.NEW_PAGE,
            "NEXT_PAGE": WD_SECTION.NEW_PAGE,
        }.get(start_type, WD_SECTION.NEW_PAGE)
        section = doc.add_section(section_type)
        page = settings.get("page") or {}
        if isinstance(page, dict):
            margins = page.get("margins") or {}
            if _is_num(margins.get("top_mm")):
                section.top_margin = Mm(margins["top_mm"])
            if _is_num(margins.get("bottom_mm")):
                section.bottom_margin = Mm(margins["bottom_mm"])
            if _is_num(margins.get("left_mm")):
                section.left_margin = Mm(margins["left_mm"])
            if _is_num(margins.get("right_mm")):
                section.right_margin = Mm(margins["right_mm"])
        for attr, part_name in (("header", "header"), ("footer", "footer")):
            content = settings.get(attr)
            if not isinstance(content, dict) or not content.get("content"):
                continue
            part = getattr(section, part_name)
            part.is_linked_to_previous = bool(content.get("linked_to_previous", False))
            paragraph = part.paragraphs[0]
            paragraph.clear()
            self._render_field_text(
                paragraph,
                str(content["content"]),
                TemplateStyle.from_dict(content.get("style")),
            )
        page_number = settings.get("page_number") or {}
        if isinstance(page_number, dict) and page_number.get("restart"):
            pg_num = OxmlElement("w:pgNumType")
            pg_num.set(qn("w:start"), str(int(page_number.get("start", 1))))
            section._sectPr.append(pg_num)

    def _render_table(self, doc: Document, item: dict,
                      template: Template) -> None:
        title = item.get("title")
        if title:
            block = self._resolve_block(template, "table_caption",
                                        "table_caption")
            style = self._block_style(block, "self")
            self._add_styled_paragraph(doc, str(title), style)
        block = self._resolve_block(template, item.get("template_key"), "table")
        table_settings = dict(block.settings or {})
        headers = item.get("headers") or table_settings.get("headers") or []
        rows = item.get("rows") or []
        n_cols = max(len(headers), len(rows[0]) if rows else 0, 1)
        n_rows = len(rows) + (1 if headers else 0)
        table = doc.add_table(rows=n_rows, cols=n_cols)
        try:
            table.style = "Table Grid"
        except (KeyError, ValueError):
            pass
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        widths = item.get("column_widths_cm") or table_settings.get("column_widths_cm") or []
        for col, width_cm in enumerate(widths):
            if not _is_num(width_cm) or col >= n_cols:
                continue
            for row in table.rows:
                row.cells[col].width = Mm(width_cm)
        r = 0
        if headers:
            for c, h in enumerate(headers):
                cell = table.cell(0, c)
                cell.text = ""
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].add_run(str(h))
            r = 1
        for row in rows:
            for c, val in enumerate(row):
                cell = table.cell(r, c)
                cell.text = ""
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].add_run(str(val))
            r += 1

    def _render_references(self, doc: Document, item: dict, template: Template,
                           meta: dict) -> None:
        block = self._resolve_block(template, "references", "references")
        title_style = self._block_style(block, "title")
        item_style = self._block_style(block, "item")
        settings = block.settings or {}
        style_name = (settings.get("style") or meta.get("reference_style")
                      or "gb7714")
        items = item.get("items") or []
        formatted = self._format_reference_items(items, style_name)
        self._add_styled_paragraph(doc, "参考文献", title_style, outline_level=1)
        hanging = settings.get("hanging_indent_pt")
        for ref in formatted:
            p = doc.add_paragraph()
            run = p.add_run(str(ref))
            self._apply_style(p, run, item_style)
            if _is_num(hanging):
                pf = p.paragraph_format
                pf.left_indent = Pt(hanging)
                pf.first_line_indent = Pt(-hanging)

    def _render_special(self, doc: Document, text: str, template: Template,
                        key: str) -> None:
        """致谢 / 附录：标题 + 内容段落。"""
        block = self._resolve_block(template, key, key)
        title_style = self._block_style(block, "title")
        content_style = self._block_style(block, "content")
        title = {"acknowledgement": "致  谢", "appendix": "附  录"}.get(
            key, key)
        self._add_styled_paragraph(doc, title, title_style, outline_level=1)
        for para in str(text).splitlines() or [str(text)]:
            if para.strip():
                self._add_styled_paragraph(doc, para, content_style)

    # ------------------------------------------------------------------
    # 编号
    # ------------------------------------------------------------------
    def _heading_number(self, level: int) -> str | None:
        """按模板 numbering 生成标题编号；返回 None 表示不编号。"""
        numbering = self._numbering
        if not numbering.get("enabled", True):
            return None
        self._counters[level - 1] += 1
        for i in range(level, 4):
            self._counters[i] = 0
        pattern = numbering.get(f"h{level}")
        if not pattern:
            pattern = _DEFAULT_NUMBERING_PATTERNS.get(level, "{n}")
        replacements = {
            "{chinese}": _cn_num(self._counters[0]),
            "{n}": str(self._counters[level - 1]),
            "{h1}": str(self._counters[0]),
            "{h2}": str(self._counters[1]),
            "{h3}": str(self._counters[2]),
            "{h4}": str(self._counters[3]),
        }
        out = pattern
        for k, v in replacements.items():
            out = out.replace(k, v)
        return out

    @staticmethod
    def _normalize_heading_text(text: str) -> str:
        """Normalize legacy ``chapter-section.subsection`` heading prefixes.

        Older content generators emit forms such as ``1-1.1 Title``. They are
        normalized before duplicate-number detection so the renderer preserves a
        single canonical number rather than prepending a second counter.
        """
        return re.sub(
            r"^(\d+)-(\d+)\.(\d+)(?=\s|$)",
            r"\1.\2.\3",
            text or "",
        )

    @staticmethod
    def _looks_numbered(text: str) -> bool:
        return bool(_NUMBERED_PREFIX.match(text or ""))

    def _toc_enabled(self, template: Template) -> bool:
        """模板（或兜底）是否有启用的目录区块。"""
        block = self._resolve_block(template, "toc", "toc")
        return block.enabled

    # ------------------------------------------------------------------
    # block 兜底
    # ------------------------------------------------------------------
    def _resolve_block(self, template: Template, key: str | None,
                       kind: str | None) -> TemplateBlock:
        for tpl in (template, self.default_template):
            if tpl is None:
                continue
            if key:
                b = tpl.get_block(key)
                if b is not None:
                    return b
            if kind:
                for b in tpl.blocks_by_kind(kind):
                    if b.key == key:
                        return b
                for b in tpl.blocks_by_kind(kind):
                    return b
        fallback = SAFE_DEFAULTS.get(kind) or SAFE_DEFAULTS.get(
            key) or SAFE_DEFAULTS["paragraph"]
        return TemplateBlock.from_dict(fallback)

    @staticmethod
    def _block_style(block: TemplateBlock, role: str) -> TemplateStyle:
        return block.styles.get(role) or block.styles.get("self")

    # ------------------------------------------------------------------
    # 统一样式应用（唯一硬编码样式落地处）
    # ------------------------------------------------------------------
    def _add_styled_paragraph(self, doc: Document, text: str,
                              style: TemplateStyle | None,
                              outline_level: int | None = None):
        p = doc.add_paragraph()
        if outline_level is not None:
            self._apply_heading_paragraph_style(p, outline_level)
        run = p.add_run(text)
        self._apply_style(p, run, style)
        if outline_level is not None:
            self._set_outline_level(p, outline_level)
        return p

    @staticmethod
    def _source_heading_style_names(source_docx: Path | str | None) -> tuple[str, ...]:
        """Return heading style names used by the uploaded source document.

        The source template can use nonstandard Word heading levels, such as
        Heading 3/4/5 for chapter/section/subsection.  Preserve the first
        occurrence order so semantic renderer levels map back to the same
        document styles instead of flattening all generated headings to Normal.
        """
        path = Path(source_docx) if source_docx else None
        if path is None or not path.is_file():
            return ()
        try:
            source = Document(str(path))
        except Exception:  # noqa: BLE001 - a malformed optional source falls back safely
            return ()
        names: list[str] = []
        for paragraph in source.paragraphs:
            name = getattr(paragraph.style, "name", "") or ""
            if not re.fullmatch(r"(?:Heading|鏍囬)\s*\d+", name, flags=re.IGNORECASE):
                continue
            if name not in names:
                names.append(name)
        return tuple(names)

    def _apply_heading_paragraph_style(self, paragraph, level: int) -> None:
        """Assign a semantic paragraph style before applying direct formatting.

        Templates frequently reference Word latent heading styles in paragraph
        XML without shipping concrete style definitions.  Such names can be
        read from source paragraphs but cannot be assigned by python-docx.  Try
        all source and built-in names first; if none is writable, create a
        stable custom paragraph style so generated headings never collapse to
        Normal.
        """
        candidates: list[str] = []
        if 1 <= level <= len(self._source_heading_styles):
            candidates.append(self._source_heading_styles[level - 1])
        candidates.extend((f"Heading {level}", f"鏍囬 {level}"))
        for style_name in candidates:
            try:
                paragraph.style = style_name
                return
            except (KeyError, ValueError):
                continue

        style_name = f"PW Heading {level}"
        doc = paragraph.part.document
        try:
            doc.styles[style_name]
        except KeyError:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        paragraph.style = style_name

    def _apply_style(self, paragraph, run, style: TemplateStyle | None) -> None:
        """把 TemplateStyle 应用到 paragraph + run（唯一样式入口）。"""
        if style is None:
            return
        # ---- 字符级 ----
        run.font.name = style.font_family_latin or DEFAULT_LATIN_FONT
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"),
                   style.font_family_east_asia or DEFAULT_EAST_ASIA_FONT)
        run.font.size = Pt(style.font_size_pt)
        run.font.bold = style.bold
        run.font.italic = style.italic
        run.font.underline = style.underline

        # ---- 段落级 ----
        pf = paragraph.paragraph_format
        pf.alignment = _ALIGN_MAP.get(style.alignment,
                                      WD_ALIGN_PARAGRAPH.JUSTIFY)
        if style.line_spacing_mode == LineSpacingMode.EXACT:
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(style.line_spacing_value)
        elif style.line_spacing_mode == LineSpacingMode.AT_LEAST:
            pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            pf.line_spacing = Pt(style.line_spacing_value)
        else:
            pf.line_spacing = style.line_spacing_value
        pf.space_before = Pt(style.space_before_pt)
        pf.space_after = Pt(style.space_after_pt)
        pf.keep_with_next = style.keep_with_next
        pf.page_break_before = style.page_break_before

        # ---- 首行缩进 ----
        if style.first_line_indent_unit == IndentUnit.CHARS and \
                style.first_line_indent_value > 0:
            self._set_first_line_chars(paragraph,
                                       style.first_line_indent_value)
        elif style.first_line_indent_value > 0:
            pf.first_line_indent = Pt(style.first_line_indent_value)

    @staticmethod
    def _set_first_line_chars(paragraph, chars: float) -> None:
        """首行缩进（字符）：写 Word OOXML w:firstLineChars（精确）。"""
        pPr = paragraph._p.get_or_add_pPr()
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        # w:firstLineChars 单位是字符的百分之一
        ind.set(qn("w:firstLineChars"), str(int(round(chars * 100))))

    @staticmethod
    def _set_outline_level(paragraph, level: int) -> None:
        """设置 w:outlineLvl，使 Word TOC 能识别标题层级。"""
        pPr = paragraph._p.get_or_add_pPr()
        ol = pPr.find(qn("w:outlineLvl"))
        if ol is None:
            ol = OxmlElement("w:outlineLvl")
            rPr = pPr.find(qn("w:rPr"))
            if rPr is not None:
                rPr.addprevious(ol)
            else:
                pPr.append(ol)
        ol.set(qn("w:val"), str(max(0, level - 1)))

    @staticmethod
    def _mark_toc_update(doc: Document) -> None:
        """设置 w:updateFields，Word 打开时自动更新目录。"""
        settings = doc.settings.element
        el = settings.find(qn("w:updateFields"))
        if el is None:
            el = OxmlElement("w:updateFields")
            settings.append(el)
        el.set(qn("w:val"), "true")

    # ------------------------------------------------------------------
    # 辅助
    # ------------------------------------------------------------------
    def _content_width_mm(self, template: Template) -> float:
        page = template.page or {}
        margins = page.get("margins") or {}
        w_mm, _ = self._page_size_mm(page.get("size") or "A4")
        left = margins.get("left_mm") or 0.0
        right = margins.get("right_mm") or 0.0
        return max(10.0, w_mm - left - right)

    @staticmethod
    def _format_reference_items(items: list, style_name: str) -> list[str]:
        """复用现有 reference.py 的内容格式化逻辑（GB/T 7714 等）。"""
        if not items:
            return []
        if isinstance(items[0], dict):
            try:
                from app.formatter.style import engine
                _, _, references = engine()
                return references.format_references(items, style_name,
                                                    "numeric")
            except Exception:
                return [str(x) for x in items]
        return [str(x) for x in items]


#: 对齐映射
_ALIGN_MAP = {
    TextAlign.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
    TextAlign.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
    TextAlign.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
    TextAlign.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _is_num(value: Any) -> bool:
    return isinstance(value, (int, float)) and not isinstance(value, bool)
