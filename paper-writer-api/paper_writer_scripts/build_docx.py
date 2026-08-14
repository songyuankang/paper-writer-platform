#!/usr/bin/env python3
"""Build a Chinese academic paper (.docx) from a JSON spec.

Two modes:

1. Default mode (no --template): applies the built-in Chinese academic format.
2. Template mode (--template school_template.docx): opens the school template
   and applies its styles, page setup, cover, TOC and structure. Generated
   content inherits the template's styles; existing 摘要/关键词/目录/参考文献/
   致谢 headings are kept and only their content is replaced.

Usage:
    python build_docx.py paper_spec.json [--out output.docx]
    python build_docx.py paper_spec.json --template school_template.docx \\
        --out output.docx [--report TemplateReport.md]

JSON spec format:
{
  "meta": {
    "title": "论文标题",
    "subtitle": "（可选）副标题",
    "author": "（可选）作者",
    "abstract": "摘要文字",
    "keywords": ["关键词1", "关键词2"],
    "abstract_page_break": true,
    "page": {"top_cm": 3.0, "bottom_cm": 2.5, "left_cm": 3.0, "right_cm": 2.5},
    "cover": {"school": "大学", "college": "学院", "major": "专业",
              "class_name": "班级", "name": "姓名", "student_id": "学号",
              "advisor": "指导教师", "date": "日期"},
    "reference_mode": "auto | user_provided | none",
    "reference_style": "gb7714 | apa | mla | chicago",
    "citation_style": "numeric | author_year"
  },
  "sections": [
    {"type": "h1", "text": "1 引言"},
    {"type": "p", "text": "正文段落"},
    {"type": "h2", "text": "1.1 研究背景"},
    {"type": "table", "title": "表1 变量描述",
     "headers": ["变量", "均值"], "rows": [["A", "1.2"], ["B", "3.4"]]},
    {"type": "figure", "path": "figures/fig1.png", "title": "图1 变化趋势"},
    {"type": "pagebreak"},
    {"type": "references", "items": ["[1] 作者. 题名[J]. 刊名, 年, 卷(期): 页码."]}
  ]
}

In template mode the template's page setup, styles, cover and TOC field are
kept; only the body content is replaced. A TemplateReport.md is written next to
the output (or to --report). Cover fields keep their placeholders when the
corresponding meta.cover value is missing. The TOC field is marked for update
so Word refreshes it on open.
"""

import argparse
import json
import re
import sys
from pathlib import Path

sys.dont_write_bytecode = True

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))
import parse_template
import references

CENTER = WD_ALIGN_PARAGRAPH.CENTER
LEFT = WD_ALIGN_PARAGRAPH.LEFT
CJK_RE = re.compile(r"[\u3400-\u4dbf\u4e00-\u9fff]")
STRUCTURE_ANCHOR_KEYS = ("摘要", "关键词", "目录", "参考文献", "致谢")
ANCHOR_TO_GROUP = {"摘要": "abstract", "关键词": "keywords",
                   "参考文献": "refs", "致谢": "ack"}


# ---------------------------------------------------------------- default mode

def set_run_font(run, cn_font="宋体", en_font="Times New Roman", size=12, bold=False):
    run.font.name = en_font
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), cn_font)


def add_para(doc, text, cn="宋体", size=12, bold=False, align=None,
             indent_pt=0, line_spacing=1.5, before=0, after=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    if indent_pt:
        pf.first_line_indent = Pt(indent_pt)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, cn_font=cn, size=size, bold=bold)
    return p


def setup_document(meta):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")

    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    page = meta.get("page", {})
    sec.top_margin = Cm(page.get("top_cm", 3.0))
    sec.bottom_margin = Cm(page.get("bottom_cm", 2.5))
    sec.left_margin = Cm(page.get("left_cm", 3.0))
    sec.right_margin = Cm(page.get("right_cm", 2.5))
    return doc


def add_title_block(doc, meta):
    if meta.get("title"):
        add_para(doc, meta["title"], cn="黑体", size=22, bold=True,
                 align=CENTER, line_spacing=1.25, before=6, after=6)
    if meta.get("subtitle"):
        add_para(doc, meta["subtitle"], cn="黑体", size=16, bold=True,
                 align=CENTER, line_spacing=1.25, after=6)
    if meta.get("author"):
        add_para(doc, meta["author"], cn="宋体", size=14, align=CENTER,
                 line_spacing=1.25, after=12)


def add_abstract(doc, meta):
    """渲染中英文摘要与关键词；默认构建器和模板渲染器保持一致。"""
    abstract = str(meta.get("abstract") or "").strip()
    keywords = meta.get("keywords") or []
    if abstract:
        add_para(doc, "摘  要", cn="黑体", size=15, bold=True,
                 align=CENTER, line_spacing=1.25, before=6, after=6)
        add_para(doc, abstract, cn="宋体", size=12,
                 indent_pt=24, line_spacing=1.5)
    if keywords:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.space_before = Pt(6)
        label = p.add_run("关键词：")
        set_run_font(label, cn_font="黑体", size=12, bold=True)
        body = p.add_run("；".join(str(item) for item in keywords))
        set_run_font(body, cn_font="宋体", size=12)

    abstract_en = str(meta.get("abstract_en") or "").strip()
    keywords_en = meta.get("keywords_en") or []
    if abstract_en:
        # 英文摘要独立成页，确保位于目录之前且不与中文摘要混排。
        if abstract or keywords:
            doc.add_page_break()
        add_para(doc, "Abstract", cn="Times New Roman", size=15, bold=True,
                 align=CENTER, line_spacing=1.25, before=6, after=6)
        add_para(doc, abstract_en, cn="Times New Roman", size=12,
                 indent_pt=0, line_spacing=1.5)
        if keywords_en:
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing = 1.5
            pf.space_before = Pt(6)
            label = p.add_run("Keywords: ")
            set_run_font(label, cn_font="Times New Roman", size=12, bold=True)
            body = p.add_run("; ".join(str(item) for item in keywords_en))
            set_run_font(body, cn_font="Times New Roman", size=12)
    if meta.get("abstract_page_break"):
        doc.add_page_break()

def add_heading(doc, text, level):
    conf = {1: (16, 12, 12), 2: (14, 6, 6), 3: (12, 6, 6)}.get(level, (12, 6, 6))
    add_para(doc, text, cn="黑体", size=conf[0], bold=True, align=LEFT,
             line_spacing=1.25, before=conf[1], after=conf[2])


def add_table(doc, item):
    if item.get("title"):
        add_para(doc, item["title"], cn="黑体", size=10.5, bold=True,
                 align=CENTER, line_spacing=1.25, before=6, after=3)
    headers = item.get("headers") or []
    rows = item.get("rows") or []
    n_cols = max(len(headers), len(rows[0]) if rows else 0, 1)
    n_rows = len(rows) + (1 if headers else 0)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    r = 0
    if headers:
        for c, h in enumerate(headers):
            cell = table.cell(0, c)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = CENTER
            set_run_font(p.add_run(str(h)), cn_font="黑体", size=10.5, bold=True)
        r = 1
    for row in rows:
        for c, val in enumerate(row):
            cell = table.cell(r, c)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = CENTER
            set_run_font(p.add_run(str(val)), cn_font="宋体", size=10.5)
        r += 1


def add_figure(doc, item, base_dir, max_width_cm=14.0):
    path = Path(item["path"])
    if not path.is_absolute():
        path = Path(base_dir) / path
    p = doc.add_paragraph()
    p.alignment = CENTER
    run = p.add_run()
    width = Cm(max_width_cm)
    try:
        from PIL import Image
        with Image.open(path) as img:
            dpi = img.info.get("dpi", (96, 96))[0] or 96
            natural_cm = img.size[0] * 2.54 / dpi
            if 3.0 < natural_cm < max_width_cm:
                width = Cm(natural_cm)
    except Exception:
        pass
    run.add_picture(str(path), width=width)
    if item.get("title"):
        add_para(doc, item["title"], cn="黑体", size=10.5, bold=True,
                 align=CENTER, line_spacing=1.25, before=3, after=6)


def add_references(doc, items):
    add_para(doc, "参考文献", cn="黑体", size=16, bold=True, align=LEFT,
             line_spacing=1.25, before=12, after=6)
    for ref in items:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Cm(0.74)
        pf.first_line_indent = Cm(-0.74)
        pf.line_spacing = 1.25
        set_run_font(p.add_run(ref), cn_font="宋体", size=10.5)


def build_default(doc, spec, meta, spec_base):
    add_title_block(doc, meta)
    add_abstract(doc, meta)
    for item in spec.get("sections", []):
        kind = item["type"]
        if kind == "h1":
            add_heading(doc, item["text"], 1)
        elif kind == "h2":
            add_heading(doc, item["text"], 2)
        elif kind == "h3":
            add_heading(doc, item["text"], 3)
        elif kind == "p":
            add_para(doc, item["text"], cn="宋体", size=12,
                     indent_pt=24, line_spacing=1.5)
        elif kind == "table":
            add_table(doc, item)
        elif kind == "figure":
            add_figure(doc, item, spec_base)
        elif kind == "pagebreak":
            doc.add_page_break()
        elif kind == "references":
            add_references(doc, item.get("items", []))
        else:
            print(f"[WARN] Unknown section type: {kind}", file=sys.stderr)


# -------------------------------------------------------------- template mode

def role_style_name(profile, role, fallback=None):
    st = profile["styles"].get(role)
    if st and st.get("name"):
        return st["name"]
    return fallback


def add_styled_para(doc, text, style_name=None):
    p = doc.add_paragraph()
    if style_name:
        try:
            p.style = doc.styles[style_name]
        except KeyError:
            pass
    p.add_run(text)
    return p


def set_outline_level(p, level):
    pPr = p._p.get_or_add_pPr()
    ol = pPr.find(qn("w:outlineLvl"))
    if ol is None:
        ol = OxmlElement("w:outlineLvl")
        rPr = pPr.find(qn("w:rPr"))
        if rPr is not None:
            rPr.addprevious(ol)
        else:
            pPr.append(ol)
    ol.set(qn("w:val"), str(level))


def add_heading_tmpl(doc, text, level, style_name):
    p = add_styled_para(doc, text, style_name)
    if p.style.name and re.match(r"^(标题|Heading)\s*\d", p.style.name or ""):
        return p
    try:
        p.style = doc.styles[f"Heading {level}"]
    except KeyError:
        conf = {1: (16, True), 2: (14, True), 3: (12, True)}.get(level, (12, True))
        for run in p.runs:
            set_run_font(run, cn_font="黑体", size=conf[0], bold=conf[1])
    set_outline_level(p, level - 1)
    return p


def add_table_tmpl(doc, item, table_style, caption_style):
    headers = item.get("headers") or []
    rows = item.get("rows") or []
    n_cols = max(len(headers), len(rows[0]) if rows else 0, 1)
    n_rows = len(rows) + (1 if headers else 0)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    try:
        table.style = doc.styles[table_style] if table_style else "Table Grid"
    except (KeyError, ValueError):
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    r = 0
    if headers:
        for c, h in enumerate(headers):
            cell = table.cell(0, c)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = CENTER
            p.add_run(str(h))
        r = 1
    for row in rows:
        for c, val in enumerate(row):
            cell = table.cell(r, c)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = CENTER
            p.add_run(str(val))
        r += 1
    return table


def make_block(doc, blk, spec_base, table_style, caption_style, max_width_cm):
    kind = blk[0]
    if kind == "para":
        return [add_styled_para(doc, blk[1], blk[2])._p]
    if kind == "heading":
        return [add_heading_tmpl(doc, blk[1], blk[2], blk[3])._p]
    if kind == "keywords":
        items = blk[1]
        p = doc.add_paragraph()
        label = p.add_run("关键词：")
        set_run_font(label, cn_font="黑体", size=12, bold=True)
        body = p.add_run("；".join(items))
        set_run_font(body, cn_font="宋体", size=12)
        return [p._p]
    if kind == "table":
        els = []
        if blk[1].get("title"):
            els.append(add_styled_para(doc, blk[1]["title"], caption_style)._p)
        els.append(add_table_tmpl(doc, blk[1], blk[2], blk[3])._tbl)
        return els
    if kind == "figure":
        path = Path(blk[1]["path"])
        if not path.is_absolute():
            path = Path(spec_base) / path
        p = doc.add_paragraph()
        p.alignment = CENTER
        run = p.add_run()
        width = Cm(max_width_cm)
        try:
            from PIL import Image
            with Image.open(path) as img:
                dpi = img.info.get("dpi", (96, 96))[0] or 96
                natural_cm = img.size[0] * 2.54 / dpi
                if 3.0 < natural_cm < max_width_cm:
                    width = Cm(natural_cm)
        except Exception:
            pass
        run.add_picture(str(path), width=width)
        els = [p._p]
        if blk[1].get("title"):
            els.append(add_styled_para(doc, blk[1]["title"], caption_style)._p)
        return els
    if kind == "pagebreak":
        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)
        return [p._p]
    if kind == "refs":
        els = []
        for ref in blk[1]:
            els.append(add_styled_para(doc, ref, blk[2])._p)
        return els
    raise ValueError(f"Unknown block kind: {kind}")


def insert_blocks(doc, blocks, anchor_el, spec_base, table_style, caption_style,
                  max_width_cm, before=False):
    if not blocks:
        return
    if anchor_el is None:
        for blk in blocks:
            make_block(doc, blk, spec_base, table_style, caption_style,
                       max_width_cm)
        return
    cursor = anchor_el
    for blk in (reversed(blocks) if before else blocks):
        els = make_block(doc, blk, spec_base, table_style, caption_style,
                         max_width_cm)
        for el in (reversed(els) if before else els):
            if before:
                cursor.addprevious(el)
            else:
                cursor.addnext(el)
            cursor = el


def find_anchor_key(text):
    n = parse_template.norm(text)
    for key in STRUCTURE_ANCHOR_KEYS:
        if n.startswith(key.casefold()):
            return key
    return None


def collect_blocks(doc, profile, spec, meta, existing_anchors):
    """Route generated content to template anchors; return block groups."""
    body_style = role_style_name(profile, "body")
    h1 = role_style_name(profile, "heading1")
    h2 = role_style_name(profile, "heading2")
    h3 = role_style_name(profile, "heading3")
    cap_style = role_style_name(profile, "caption")
    ref_style = role_style_name(profile, "reference") or body_style
    table_style = profile["styles"].get("table", {}).get("name") if \
        profile["styles"].get("table") else None
    page = profile["page"]
    content_width = max(8.0, (page.get("page_width_cm") or 21.0)
                        - (page.get("left_margin_cm") or 3.0)
                        - (page.get("right_margin_cm") or 3.0))

    groups = {"abstract": [], "keywords": [], "body": [], "refs": [], "ack": []}
    current = "body"
    if meta.get("abstract"):
        if "摘要" in existing_anchors:
            groups["abstract"].append(("para", meta["abstract"], body_style))
        else:
            groups["abstract"].append(("heading", "摘要", 1, h1))
            groups["abstract"].append(("para", meta["abstract"], body_style))
    if meta.get("keywords"):
        if "关键词" in existing_anchors:
            groups["keywords"].append(("keywords", meta["keywords"], body_style))
        else:
            groups["abstract"].append(("keywords", meta["keywords"], body_style))

    for item in spec.get("sections", []):
        kind = item["type"]
        if kind == "h1":
            key = find_anchor_key(item["text"])
            if key in existing_anchors and key in ANCHOR_TO_GROUP:
                current = ANCHOR_TO_GROUP[key]
                continue
            current = "body"
            groups["body"].append(("heading", item["text"], 1, h1))
        elif kind == "h2":
            groups[current].append(("heading", item["text"], 2, h2))
        elif kind == "h3":
            groups[current].append(("heading", item["text"], 3, h3))
        elif kind == "p":
            groups[current].append(("para", item["text"], body_style))
        elif kind == "table":
            groups[current].append(("table", item, table_style, cap_style))
        elif kind == "figure":
            groups[current].append(("figure", item, None, None))
        elif kind == "pagebreak":
            groups[current].append(("pagebreak",))
        elif kind == "references":
            items = item.get("items", [])
            if "参考文献" in existing_anchors:
                groups["refs"].append(("refs", items, ref_style))
            else:
                groups["body"].append(("heading", "参考文献", 1, h1))
                groups["body"].append(("refs", items, ref_style))
        else:
            print(f"[WARN] Unknown section type: {kind}", file=sys.stderr)
    return groups, {"body_style": body_style, "table_style": table_style,
                    "caption_style": cap_style, "max_width_cm": content_width}


def find_anchors(doc, profile):
    anchors = {}
    toc_indices = set(profile["toc"].get("field_paragraphs", []))
    keep = set(toc_indices)
    for item in profile["structure"].get("items", []):
        idx = item["index"]
        if item["kind"] == "cover":
            keep.add(idx)
        elif item["kind"] == "heading":
            key = find_anchor_key(item["text"])
            if key:
                anchors[key] = doc.paragraphs[idx]
                keep.add(idx)
    return anchors, keep


def cover_values(meta):
    cover = meta.get("cover") or {}
    return {
        "school": cover.get("school"), "college": cover.get("college"),
        "major": cover.get("major"), "class_name": cover.get("class_name"),
        "name": cover.get("name"), "student_id": cover.get("student_id"),
        "advisor": cover.get("advisor"), "date": cover.get("date"),
        "grade": cover.get("grade"), "title": cover.get("title") or meta.get("title"),
    }


def fill_paragraph(p, value):
    text = p.text
    m = re.search(r"[＿_]{2,}", text)
    if m:
        start, end = m.start(), m.end()
        pos = 0
        for run in p.runs:
            run_start, run_end = pos, pos + len(run.text)
            if run_start <= start < run_end:
                run.text = (run.text[:start - run_start] + value
                            + run.text[end - run_start:])
                return True
            pos = run_end
        return False
    if re.search(r"[：:]\s*$", text):
        if p.runs:
            p.runs[-1].text = p.runs[-1].text.rstrip() + value
            return True
    return False


def fill_cover(doc, profile, meta):
    values = cover_values(meta)
    filled, placeholders = [], []
    for fld in profile["cover"].get("fields", []):
        idx = fld["paragraph_index"]
        if idx >= len(doc.paragraphs):
            continue
        p = doc.paragraphs[idx]
        value = values.get(fld["key"])
        if value and fill_paragraph(p, str(value)):
            filled.append(fld["keyword"])
        else:
            placeholders.append(fld["keyword"])
    return {"filled": filled, "placeholders": placeholders}


def mark_toc_update(doc):
    settings = doc.settings.element
    el = settings.find(qn("w:updateFields"))
    if el is None:
        el = OxmlElement("w:updateFields")
        settings.append(el)
    el.set(qn("w:val"), "true")


def build_with_template(doc, profile, spec, meta, spec_base):
    anchors, keep = find_anchors(doc, profile)
    cover_status = fill_cover(doc, profile, meta)

    first_structural = None
    for item in profile["structure"].get("items", []):
        if item["kind"] in ("heading", "toc"):
            first_structural = item["index"]
            break

    body = doc.element.body
    first_structural_el = (doc.paragraphs[first_structural]._p
                           if first_structural is not None else None)
    for i, p in enumerate(doc.paragraphs):
        if i not in keep:
            p._p.getparent().remove(p._p)
    for tbl in doc.tables:
        keep_table = False
        if first_structural_el is not None:
            try:
                keep_table = body.index(tbl._tbl) < body.index(first_structural_el)
            except ValueError:
                keep_table = False
        if not keep_table:
            tbl._tbl.getparent().remove(tbl._tbl)

    groups, style_info = collect_blocks(doc, profile, spec, meta, set(anchors))
    cap_style = style_info["caption_style"]
    max_width = style_info["max_width_cm"]

    body_anchor = (anchors.get("参考文献") or anchors.get("致谢"))._p \
        if (anchors.get("参考文献") or anchors.get("致谢")) else None
    abstract_anchor = anchors["摘要"]._p if "摘要" in anchors else body_anchor
    keywords_anchor = anchors["关键词"]._p if "关键词" in anchors else None
    refs_anchor = anchors["参考文献"]._p if "参考文献" in anchors else body_anchor
    ack_anchor = anchors["致谢"]._p if "致谢" in anchors else None

    if "关键词" in anchors:
        kw_anchor = anchors["关键词"]
        if "：" in kw_anchor.text or ":" in kw_anchor.text:
            value = "关键词：" + "；".join(meta.get("keywords") or [])
            if kw_anchor.runs:
                kw_anchor.runs[0].text = value
                for extra in kw_anchor.runs[1:]:
                    extra.text = ""
            keywords_anchor = None

    insert_blocks(doc, groups["abstract"], abstract_anchor, spec_base,
                  style_info["table_style"], cap_style, max_width)
    if "关键词" in anchors:
        if keywords_anchor is not None:
            insert_blocks(doc, groups["keywords"], keywords_anchor, spec_base,
                          style_info["table_style"], cap_style, max_width)
    insert_blocks(doc, groups["body"], body_anchor, spec_base,
                  style_info["table_style"], cap_style, max_width, before=True)
    insert_blocks(doc, groups["refs"], refs_anchor, spec_base,
                  style_info["table_style"], cap_style, max_width)
    insert_blocks(doc, groups["ack"], ack_anchor, spec_base,
                  style_info["table_style"], cap_style, max_width)

    mark_toc_update(doc)
    return cover_status


# ------------------------------------------------------------------ word count

def count_cjk(text):
    return len(CJK_RE.findall(text))


def report_counts(doc):
    hanzi = 0
    non_ws = 0
    for p in doc.paragraphs:
        t = p.text
        hanzi += count_cjk(t)
        non_ws += len(t.replace(" ", "").replace("\u3000", ""))
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                t = cell.text
                hanzi += count_cjk(t)
                non_ws += len(t.replace(" ", "").replace("\u3000", ""))
    summary = {
        "hanzi": hanzi,
        "non_whitespace_chars": non_ws,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
    }
    print(json.dumps(summary, ensure_ascii=False))
    return summary


def main():
    sys.stdout.reconfigure(encoding="utf-8")
    parser = argparse.ArgumentParser(description="Build a Chinese academic paper .docx")
    parser.add_argument("spec", help="Path to paper_spec.json")
    parser.add_argument("--out", help="Output .docx path (default: spec name + .docx)")
    parser.add_argument("--template", help="School template .docx (optional)")
    parser.add_argument("--report", help="TemplateReport.md path (template mode)")
    parser.add_argument("--references", help="references.json (optional)")
    parser.add_argument("--reference-style", choices=("gb7714", "apa", "mla", "chicago"),
                        default="gb7714")
    parser.add_argument("--citation-style", choices=("numeric", "author_year"),
                        default="numeric")
    parser.add_argument("--no-reference-check", action="store_true",
                        help="Skip writing ReferenceCheck.md")
    args = parser.parse_args()

    spec_path = Path(args.spec)
    spec = json.loads(spec_path.read_text(encoding="utf-8"))
    meta = spec.get("meta", {})
    out = Path(args.out) if args.out else spec_path.with_suffix(".docx")
    out.parent.mkdir(parents=True, exist_ok=True)

    ref_stats = None
    if args.references:
        entries = references.load_entries(args.references)
        mode = meta.get("reference_mode", "auto")
        if mode != "none":
            has_ref_section = any(
                item.get("type") == "references" for item in spec.get("sections", []))
            if not has_ref_section:
                formatted = references.format_references(
                    entries, args.reference_style, args.citation_style)
                spec.setdefault("sections", []).append(
                    {"type": "references", "items": formatted})
        ref_stats = references.check_references(
            spec, entries, args.reference_style, args.citation_style)
        if not args.no_reference_check:
            ref_report = out.with_name("ReferenceCheck.md")
            ref_report.write_text(
                references.render_check(ref_stats, args.reference_style,
                                        args.citation_style),
                encoding="utf-8")

    if args.template:
        if not Path(args.template).exists():
            print(f"[ERROR] Template not found: {args.template}", file=sys.stderr)
            sys.exit(1)
        profile = parse_template.parse_document(args.template)
        doc = Document(args.template)
        cover_status = build_with_template(doc, profile, spec, meta, spec_path.parent)
        doc.save(out)
        report = Path(args.report) if args.report else out.with_name("TemplateReport.md")
        report.write_text(parse_template.render_report(profile), encoding="utf-8")
        comp = profile["compatibility"]
        print(json.dumps({
            "mode": "template",
            "template": args.template,
            "compatibility_score": comp["score"],
            "missing": comp["missing"],
            "cover": cover_status,
            "toc_update_marked": True,
            "report": str(report),
            "reference_check": ref_stats,
        }, ensure_ascii=False, indent=2))
    else:
        doc = setup_document(meta)
        build_default(doc, spec, meta, spec_path.parent)
        doc.save(out)
        print(f"[OK] Saved: {out}")
        print(json.dumps({"mode": "default",
                          "reference_check": ref_stats}, ensure_ascii=False))

    report_counts(doc)


if __name__ == "__main__":
    main()
