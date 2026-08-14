#!/usr/bin/env python3
"""Parse a school thesis template (.docx) into a machine-readable profile JSON
and a human-readable TemplateReport.md (including a compatibility score).

Priority is given to Word Styles (styles.xml) and document defaults rather than
analyzing plain text. Structure (cover, TOC, headings such as 摘要/关键词/
参考文献/致谢) is detected by paragraph styles, TOC fields, and heading text.

Usage:
    python parse_template.py school_template.docx \\
        --out template_profile.json --report TemplateReport.md

The profile JSON is also imported by build_docx.py when --template is used, so
generated content inherits the template's styles.
"""

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

STRUCT_KEYS = [
    "摘要", "关键词", "目录", "参考文献", "致谢",
    "Abstract", "Keywords", "Contents", "References", "Acknowledg",
]

COVER_FIELDS = {
    "论文题目": "title", "题目": "title", "学校": "school", "学院": "college",
    "专业": "major", "班级": "class_name", "姓名": "name", "学号": "student_id",
    "指导教师": "advisor", "导师": "advisor", "日期": "date", "年级": "grade",
}

STYLE_CANDIDATES = {
    "heading1": ["标题 1", "Heading 1", "标题一", "一级标题"],
    "heading2": ["标题 2", "Heading 2", "标题二", "二级标题"],
    "heading3": ["标题 3", "Heading 3", "标题三", "三级标题"],
    "body": ["正文", "Normal", "正文文本"],
    "caption": ["题注", "Caption"],
    "reference": ["参考文献", "Bibliography", "引用"],
}


def norm(text):
    """Strip all whitespace and casefold for fuzzy matching."""
    return re.sub(r"[\s\u3000]+", "", text or "").casefold()


def emu_to_cm(value):
    return round(value / 360000.0, 2) if value is not None else None


def half_pt_to_pt(value):
    return round(value / 2.0, 1) if value is not None else None


def read_outline_level(element):
    """Read w:outlineLvl from a style or paragraph element, or None."""
    if element is None:
        return None
    pPr = element.find(qn("w:pPr"))
    if pPr is None:
        return None
    ol = pPr.find(qn("w:outlineLvl"))
    if ol is None:
        return None
    try:
        return int(ol.get(qn("w:val")))
    except (TypeError, ValueError):
        return None


def find_style(doc, candidates):
    """Return the first defined style matching any candidate name, or None."""
    known = {s.name for s in doc.styles}
    for cand in candidates:
        if cand in known:
            try:
                return doc.styles[cand]
            except KeyError:
                continue
    return None


def style_font_info(style):
    if style is None:
        return None
    info = {"name": style.name}
    info["ascii"] = style.font.name
    info["east_asia"] = None
    info["size_pt"] = style.font.size.pt if style.font.size else None
    info["bold"] = style.font.bold
    if style.font.color and style.font.color.rgb:
        info["color"] = str(style.font.color.rgb)
    rpr = style.element.rPr
    if rpr is not None:
        rfonts = rpr.rFonts
        if rfonts is not None:
            info["east_asia"] = rfonts.get(qn("w:eastAsia"))
    return info


def style_para_info(style):
    if style is None:
        return None
    pf = style.paragraph_format
    info = {
        "alignment": str(pf.alignment) if pf.alignment is not None else None,
        "line_spacing": (pf.line_spacing.pt if hasattr(pf.line_spacing, "pt")
                         else pf.line_spacing),
        "space_before_pt": pf.space_before.pt if pf.space_before else None,
        "space_after_pt": pf.space_after.pt if pf.space_after else None,
        "first_line_indent_pt": (pf.first_line_indent.pt
                                 if pf.first_line_indent else None),
        "outline_level": read_outline_level(style.element),
    }
    return info


def paragraph_style_info(paragraph):
    """Return the effective direct formatting of one representative paragraph.

    Many school templates put every paragraph in ``Normal`` and apply the real
    typography directly to runs/paragraphs.  Reading only styles.xml therefore
    loses the visible template format.
    """
    run = next((r for r in paragraph.runs if r.text), None)
    font = {"name": paragraph.style.name if paragraph.style else None,
            "ascii": run.font.name if run else None,
            "east_asia": None,
            "size_pt": run.font.size.pt if run and run.font.size else None,
            "bold": run.bold if run else None}
    if run is not None and run._element.rPr is not None:
        rfonts = run._element.rPr.rFonts
        if rfonts is not None:
            font["east_asia"] = rfonts.get(qn("w:eastAsia"))
    pf = paragraph.paragraph_format
    para = {
        "alignment": str(paragraph.alignment) if paragraph.alignment is not None else None,
        "line_spacing": (pf.line_spacing.pt if hasattr(pf.line_spacing, "pt")
                         else pf.line_spacing),
        "space_before_pt": pf.space_before.pt if pf.space_before else None,
        "space_after_pt": pf.space_after.pt if pf.space_after else None,
        "first_line_indent_pt": (pf.first_line_indent.pt
                                 if pf.first_line_indent else None),
        "outline_level": read_outline_level(paragraph._p),
    }
    return {"name": font["name"], "font": font, "paragraph": para}


def doc_defaults(doc):
    """Read w:docDefaults (rPrDefault/pPrDefault) from styles.xml."""
    out = {"font": {}, "paragraph": {}}
    styles_el = doc.styles.element
    dd = styles_el.find(qn("w:docDefaults"))
    if dd is None:
        return out
    rpr = dd.find(qn("w:rPrDefault"))
    if rpr is not None:
        rpr = rpr.find(qn("w:rPr"))
        if rpr is not None:
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is not None:
                out["font"]["ascii"] = rfonts.get(qn("w:ascii"))
                out["font"]["east_asia"] = rfonts.get(qn("w:eastAsia"))
                out["font"]["h_ansi"] = rfonts.get(qn("w:hAnsi"))
            sz = rpr.find(qn("w:sz"))
            if sz is not None:
                out["font"]["size_pt"] = half_pt_to_pt(int(sz.get(qn("w:val"))))
    ppr = dd.find(qn("w:pPrDefault"))
    if ppr is not None:
        ppr = ppr.find(qn("w:pPr"))
        if ppr is not None:
            spacing = ppr.find(qn("w:spacing"))
            if spacing is not None:
                if spacing.get(qn("w:line")):
                    out["paragraph"]["line"] = int(spacing.get(qn("w:line")))
                    out["paragraph"]["line_rule"] = spacing.get(qn("w:lineRule"))
                if spacing.get(qn("w:before")):
                    out["paragraph"]["before_twips"] = int(spacing.get(qn("w:before")))
                if spacing.get(qn("w:after")):
                    out["paragraph"]["after_twips"] = int(spacing.get(qn("w:after")))
            ind = ppr.find(qn("w:ind"))
            if ind is not None:
                if ind.get(qn("w:firstLine")):
                    out["paragraph"]["first_line_twips"] = int(ind.get(qn("w:firstLine")))
                if ind.get(qn("w:firstLineChars")):
                    out["paragraph"]["first_line_chars"] = int(ind.get(qn("w:firstLineChars")))
    return out


def page_info(section):
    info = {
        "page_width_cm": emu_to_cm(section.page_width),
        "page_height_cm": emu_to_cm(section.page_height),
        "top_margin_cm": emu_to_cm(section.top_margin),
        "bottom_margin_cm": emu_to_cm(section.bottom_margin),
        "left_margin_cm": emu_to_cm(section.left_margin),
        "right_margin_cm": emu_to_cm(section.right_margin),
        "header_distance_cm": emu_to_cm(section.header_distance),
        "footer_distance_cm": emu_to_cm(section.footer_distance),
        "page_number_format": "decimal",
        "page_number_start": None,
        "gutter_cm": None,
        "header": {"has_content": False, "different_first_page": False},
        "footer": {"has_content": False, "different_first_page": False},
    }
    sect_pr = section._sectPr
    pg_mar = sect_pr.find(qn("w:pgMar"))
    if pg_mar is not None and pg_mar.get(qn("w:gutter")):
        info["gutter_cm"] = round(int(pg_mar.get(qn("w:gutter"))) / 567.0, 2)
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is not None:
        if pg_num.get(qn("w:fmt")):
            info["page_number_format"] = pg_num.get(qn("w:fmt"))
        if pg_num.get(qn("w:start")):
            info["page_number_start"] = int(pg_num.get(qn("w:start")))
    info["header"]["different_first_page"] = bool(
        sect_pr.find(qn("w:titlePg")) is not None)
    try:
        info["header"]["has_content"] = any(
            (p.text.strip() or p._p.xpath(".//w:drawing")
             or p._p.xpath(".//w:fldSimple")) for p in section.header.paragraphs)
        info["footer"]["has_content"] = any(
            (p.text.strip() or p._p.xpath(".//w:drawing")
             or p._p.xpath(".//w:fldSimple")) for p in section.footer.paragraphs)
    except Exception:
        pass
    return info


def is_toc_paragraph(p):
    for el in p._p.xpath(".//w:instrText"):
        instr = el.text or ""
        if "TOC" in instr.upper():
            return True
    for el in p._p.xpath(".//w:fldSimple"):
        instr = el.get(qn("w:instr")) or ""
        if "TOC" in instr.upper():
            return True
    return False


def style_heading_level(p):
    style = p.style
    if style is None:
        return None
    name = style.name or ""
    m = re.search(r"(标题|Heading)\s*([1-9])", name, re.I)
    if m:
        return int(m.group(2))
    ol = read_outline_level(p._p) or read_outline_level(style.element)
    if ol is not None and 0 <= ol <= 8:
        return ol + 1
    return None


def structural_key(text):
    n = norm(text)
    for key in STRUCT_KEYS:
        if key.casefold() in n:
            return key
    return None


def analyze_structure(doc):
    """Classify every paragraph as cover / toc / heading / body."""
    items = []
    first_structural = None
    for i, p in enumerate(doc.paragraphs):
        if is_toc_paragraph(p):
            kind = "toc"
        else:
            level = style_heading_level(p)
            key = structural_key(p.text)
            if level is not None or (key and key in ("摘要", "关键词", "目录",
                                                     "参考文献", "致谢", "Abstract",
                                                     "Keywords", "Contents",
                                                     "References", "Acknowledg")):
                kind = "heading"
            else:
                kind = "body"
        items.append({"index": i, "kind": kind, "text": p.text.strip()[:80],
                      "style": p.style.name if p.style else None})
        if first_structural is None:
            if kind == "toc" or (kind == "heading" and structural_key(p.text)):
                first_structural = i
    for item in items:
        if first_structural is not None and item["index"] < first_structural:
            item["kind"] = "cover"
    return items, first_structural


def detect_cover_fields(doc, cover_indices):
    fields = []
    for idx in cover_indices:
        if idx >= len(doc.paragraphs):
            continue
        p = doc.paragraphs[idx]
        n = norm(p.text)
        if not n:
            continue
        for keyword, key in COVER_FIELDS.items():
            if keyword.casefold() in n:
                fields.append({"keyword": keyword, "key": key,
                               "paragraph_index": idx,
                               "text": p.text.strip()[:80]})
                break
    return fields


def detect_toc_styles(doc):
    names = [s.name for s in doc.styles if s.name]
    return [n for n in names if re.match(r"^(目录|TOC)\s*\d", n, re.I)]


def compute_compatibility(profile):
    checks = [
        ("页面设置（纸张/页边距）", 10,
         all(profile["page"][k] is not None for k in
             ("page_width_cm", "page_height_cm", "top_margin_cm", "left_margin_cm"))),
        ("文档默认字体（docDefaults）", 10,
         bool(profile["fonts"]["default"].get("east_asia") or
              profile["fonts"]["default"].get("ascii"))),
        ("正文样式", 10, profile["styles"]["body"] is not None),
        ("一级标题样式", 10, profile["styles"]["heading1"] is not None),
        ("二级标题样式", 8, profile["styles"]["heading2"] is not None),
        ("三级标题样式", 6, profile["styles"]["heading3"] is not None),
        ("题注样式", 8, profile["styles"]["caption"] is not None),
        ("参考文献样式/结构", 6,
         profile["styles"]["reference"] is not None or profile["structure"]["has_references"]),
        ("目录域（TOC field）", 10, profile["toc"]["detected"]),
        ("目录条目样式", 4, bool(profile["styles"]["toc_entries"])),
        ("封面结构", 8, profile["cover"]["detected"]),
        ("封面字段数量≥6", 2, len(profile["cover"]["fields"]) >= 6),
        ("表格样式", 5, profile["styles"]["table"] is not None),
        ("页眉页脚", 3,
         profile["page"]["header"]["has_content"] or profile["page"]["footer"]["has_content"]),
    ]
    total = sum(weight for _, weight, _ in checks)
    score = sum(weight for _, weight, ok in checks if ok)
    missing = [name for name, _, ok in checks if not ok]
    return {"score": score, "max": total, "missing": missing}


def extract_table_layouts(doc):
    """Extract reusable table layout metadata without treating cell values as data.

    The first row is retained only as a column-header hint.  All remaining
    source cell content is deliberately excluded so later paper generation
    cannot accidentally copy or fabricate template data.
    """
    layouts = []
    for ordinal, table in enumerate(doc.tables, start=1):
        rows = list(table.rows)
        first_row = rows[0].cells if rows else []
        headers = [(cell.text or "").strip() for cell in first_row]
        widths_cm = []
        for cell in first_row:
            width = getattr(cell, "width", None)
            widths_cm.append(round(width / 360000, 3) if width else None)
        layouts.append({
            "key": f"table_{ordinal}",
            "ordinal": ordinal,
            "row_count": len(rows),
            "column_count": len(first_row),
            "headers": headers,
            "column_widths_cm": widths_cm,
            "style": getattr(getattr(table, "style", None), "name", None),
        })
    return layouts


def extract_section_layouts(doc):
    """Extract section-level page, header and footer settings from DOCX."""
    layouts = []
    for ordinal, section in enumerate(doc.sections, start=1):
        page = page_info(section)
        header = "\n".join(
            p.text.strip() for p in section.header.paragraphs if p.text.strip()
        )
        footer = "\n".join(
            p.text.strip() for p in section.footer.paragraphs if p.text.strip()
        )
        layouts.append({
            "key": f"section_{ordinal}",
            "ordinal": ordinal,
            "start_type": getattr(section.start_type, "name", str(section.start_type)),
            "page": page,
            "header": {"content": header, "linked_to_previous": bool(section.header.is_linked_to_previous)},
            "footer": {"content": footer, "linked_to_previous": bool(section.footer.is_linked_to_previous)},
        })
    return layouts


def parse_document(path):
    """Parse a template docx into a profile dict (JSON-serializable)."""
    doc = Document(str(path))
    section = doc.sections[0] if doc.sections else None
    profile = {
        "source": str(path),
        "page": page_info(section) if section else {},
        "fonts": {"default": doc_defaults(doc)["font"]},
        "styles": {},
        "cover": {"detected": False, "fields": []},
        "toc": {"detected": False, "field_paragraphs": []},
        "tables": extract_table_layouts(doc),
        "sections": extract_section_layouts(doc),
        "structure": {"items": [], "has_references": False},
        "compatibility": {},
    }

    for role, candidates in STYLE_CANDIDATES.items():
        style = find_style(doc, candidates)
        profile["styles"][role] = {
            "name": style.name,
            "font": style_font_info(style),
            "paragraph": style_para_info(style),
        } if style else None

    # Keep direct-format samples too.  This is deliberately separate from
    # named-style extraction so callers can prefer visible formatting without
    # losing the source style identity.
    nonempty = [p for p in doc.paragraphs if p.text.strip()]
    if nonempty:
        profile["styles"]["title_zh_direct"] = paragraph_style_info(nonempty[0])
    body_sample = next((p for p in nonempty if len(p.text.strip()) >= 80), None)
    if body_sample is not None:
        profile["styles"]["body_direct"] = paragraph_style_info(body_sample)

    profile["styles"]["toc_entries"] = detect_toc_styles(doc)
    profile["styles"]["table"] = None
    known = {s.name for s in doc.styles}
    for cand in ("Table Grid", "网格型", "表格主题"):
        if cand in known:
            profile["styles"]["table"] = {"name": cand}
            break

    items, first_structural = analyze_structure(doc)
    toc_paras = [i["index"] for i in items if i["kind"] == "toc"]
    cover_indices = [i["index"] for i in items if i["kind"] == "cover"]
    profile["toc"] = {"detected": bool(toc_paras), "field_paragraphs": toc_paras}
    profile["cover"] = {
        "detected": bool(cover_indices),
        "fields": detect_cover_fields(doc, cover_indices),
    }
    profile["structure"] = {
        "items": items,
        "has_references": any(
            norm(i.get("text", "")).find("参考文献") >= 0 or
            norm(i.get("text", "")).find("references") >= 0
            for i in items if i["kind"] == "heading"),
    }
    profile["compatibility"] = compute_compatibility(profile)
    return profile


def render_report(profile):
    """Render TemplateReport.md from a profile."""
    p = profile["page"]
    f = profile["fonts"]["default"]
    s = profile["styles"]
    lines = [
        "# 模板解析报告（TemplateReport）",
        "",
        f"- 来源模板：`{profile['source']}`",
        f"- 兼容度评分：**{profile['compatibility']['score']} / {profile['compatibility']['max']}**",
        "",
        "## 页面设置",
        "",
        f"- 纸张：{p.get('page_width_cm')} × {p.get('page_height_cm')} cm"
        if p.get("page_width_cm") else "- 纸张：未检测到",
        f"- 页边距：上 {p.get('top_margin_cm')} / 下 {p.get('bottom_margin_cm')} / "
        f"左 {p.get('left_margin_cm')} / 右 {p.get('right_margin_cm')} cm",
        f"- 装订线：{p.get('gutter_cm') or 0} cm",
        f"- 页眉距：{p.get('header_distance_cm')} cm；页脚距：{p.get('footer_distance_cm')} cm",
        f"- 页码格式：{p.get('page_number_format')}"
        + (f"（起始 {p.get('page_number_start')}）" if p.get("page_number_start") else ""),
        f"- 页眉：{'有内容' if p.get('header', {}).get('has_content') else '无'}；"
        f"页脚：{'有内容' if p.get('footer', {}).get('has_content') else '无'}；"
        f"首页不同：{'是' if p.get('header', {}).get('different_first_page') else '否'}",
        "",
        "## 字体",
        "",
        f"- 文档默认：中文 {f.get('east_asia') or '未设置'}，西文 {f.get('ascii') or '未设置'}，"
        f"字号 {f.get('size_pt') or '未设置'} pt",
    ]

    lines += ["", "## 标题样式", ""]
    for role, label in (("heading1", "一级标题"), ("heading2", "二级标题"),
                        ("heading3", "三级标题")):
        st = s.get(role)
        if st:
            ft = st["font"] or {}
            pp = st["paragraph"] or {}
            lines.append(f"- {label}（样式 `{st['name']}`）："
                         f"字体 {ft.get('east_asia') or ft.get('ascii') or '继承默认'}，"
                         f"{ft.get('size_pt') or '继承'} pt，"
                         f"加粗 {ft.get('bold')}，行距 {pp.get('line_spacing')}，"
                         f"段前 {pp.get('space_before_pt')} pt，段后 {pp.get('space_after_pt')} pt，"
                         f"首行缩进 {pp.get('first_line_indent_pt')} pt")
        else:
            lines.append(f"- {label}：未定义样式（生成时将用内置标题样式替代）")

    body = s.get("body")
    lines += ["", "## 正文样式", ""]
    if body:
        ft = body["font"] or {}
        pp = body["paragraph"] or {}
        lines.append(f"- 样式 `{body['name']}`：字体 {ft.get('east_asia') or ft.get('ascii')}，"
                     f"{ft.get('size_pt')} pt，行距 {pp.get('line_spacing')}，"
                     f"首行缩进 {pp.get('first_line_indent_pt')} pt")
    else:
        lines.append("- 未定义独立正文样式（使用 Normal）")

    lines += ["", "## 图表样式", ""]
    cap = s.get("caption")
    lines.append(f"- 题注样式：{cap['name'] if cap else '未定义'}")
    lines.append(f"- 表格样式：{s.get('table', {}).get('name') if s.get('table') else '未定义（生成时用 Table Grid）'}")

    ref = s.get("reference")
    lines += ["", "## 参考文献样式", ""]
    lines.append(f"- 参考文献样式：{ref['name'] if ref else '未定义'}"
                 f"{'（模板含参考文献章节）' if profile['structure']['has_references'] else ''}")

    toc_entries = s.get("toc_entries") or []
    lines += ["", "## 目录样式", ""]
    lines.append(f"- 目录域（TOC field）：{'检测到' if profile['toc']['detected'] else '未检测到'}")
    lines.append(f"- 目录条目样式：{('、'.join(toc_entries)) if toc_entries else '未检测到'}")

    lines += ["", "## 封面", ""]
    if profile["cover"]["fields"]:
        lines.append(f"- 检测到封面，识别字段：{'、'.join(f['keyword'] for f in profile['cover']['fields'])}")
    else:
        lines.append("- 未检测到封面字段")

    lines += ["", "## 兼容度", ""]
    comp = profile["compatibility"]
    if comp["missing"]:
        lines.append(f"- 缺失/待补齐项：{'、'.join(comp['missing'])}")
    else:
        lines.append("- 各项齐全，可完整套用模板样式")
    return "\n".join(lines) + "\n"


def main():
    sys.stdout.reconfigure(encoding="utf-8")
    parser = argparse.ArgumentParser(description="Parse a school thesis template .docx")
    parser.add_argument("template", help="Path to the school template .docx")
    parser.add_argument("--out", default="template_profile.json",
                        help="Output profile JSON path")
    parser.add_argument("--report", default="TemplateReport.md",
                        help="Output TemplateReport.md path")
    args = parser.parse_args()

    if not Path(args.template).exists():
        print(f"[ERROR] Template not found: {args.template}", file=sys.stderr)
        sys.exit(1)

    profile = parse_document(args.template)
    Path(args.out).parent.mkdir(parents=True, exist_ok=True)
    Path(args.out).write_text(json.dumps(profile, ensure_ascii=False, indent=2),
                              encoding="utf-8")
    Path(args.report).write_text(render_report(profile), encoding="utf-8")
    comp = profile["compatibility"]
    print(json.dumps({
        "profile": args.out,
        "report": args.report,
        "compatibility_score": comp["score"],
        "max": comp["max"],
        "missing": comp["missing"],
        "cover_fields": [f["keyword"] for f in profile["cover"]["fields"]],
        "toc_detected": profile["toc"]["detected"],
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
