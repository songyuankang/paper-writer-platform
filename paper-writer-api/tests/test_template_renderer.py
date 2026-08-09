"""TemplateRenderer 单元测试。

覆盖：TemplateStyle → DOCX 全字段、页面/页边距/方向、页眉页脚 + PAGE 域、
TOC 域、heading 1-4 + outlineLvl、编号（启用/关闭/不重复）、段落、摘要、
关键词、图表题注、参考文献、缺失 block 兜底、三个真实 v2 模板渲染。

运行：``python -m unittest tests.test_template_renderer -v``
"""

from __future__ import annotations

import base64
import sys
import tempfile
import unittest
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from docx.enum.section import WD_ORIENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

from app.formatter.template import (  # noqa: E402
    DEFAULT_TEMPLATES_ROOT,
    Template,
    TemplateRenderer,
    TemplateStyle,
)
from app.formatter.template.loader import TemplateLoader  # noqa: E402
try:  # noqa: E402
    from test_template_system import make_template_dict  # noqa: E402
except ImportError:  # noqa: E402
    from tests.test_template_system import make_template_dict  # noqa: E402

#: 1x1 PNG（用于 figure 渲染测试）
_TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg==")


def make_spec(**overrides) -> dict:
    """标准论文内容 spec（与现有 paper_spec 结构一致）。"""
    spec = {
        "meta": {
            "title": "论文标题",
            "abstract": "本文研究摘要内容。",
            "keywords": ["关键词一", "关键词二"],
            "reference_style": "gb7714",
        },
        "sections": [
            {"type": "h1", "text": "引言"},
            {"type": "p", "text": "第一段正文内容。"},
            {"type": "h2", "text": "研究背景"},
            {"type": "p", "text": "第二段正文内容。"},
            {"type": "h3", "text": "子问题"},
            {"type": "h1", "text": "文献综述"},
            {"type": "references", "items": [
                "[1] 张三. 示例研究[J]. 期刊, 2025, 10(2): 1-5."]},
        ],
        "references": ["[1] 张三. 示例研究[J]. 期刊, 2025, 10(2): 1-5."],
    }
    spec.update(overrides)
    return spec


def make_template(**overrides) -> Template:
    data = make_template_dict()
    data.update(overrides)
    return Template.from_dict(data, template_id="test-tpl")


def load_basic(stem: str) -> Template:
    loader = TemplateLoader(DEFAULT_TEMPLATES_ROOT)
    return loader.load_template(DEFAULT_TEMPLATES_ROOT / "basic" / f"{stem}.json")


# ---- OOXML 读取助手 ----
def _east_asia(run):
    rpr = run._element.rPr
    if rpr is None or rpr.rFonts is None:
        return None
    return rpr.rFonts.get(qn("w:eastAsia"))


def _first_line_chars(p):
    pPr = p._p.pPr
    if pPr is None:
        return None
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        return None
    return ind.get(qn("w:firstLineChars"))


def _outline_level(p):
    pPr = p._p.pPr
    if pPr is None:
        return None
    ol = pPr.find(qn("w:outlineLvl"))
    return ol.get(qn("w:val")) if ol is not None else None


def _instr_texts(p):
    return [el.text or "" for el in p._p.xpath(".//w:instrText")]


def _update_fields(doc):
    el = doc.settings.element.find(qn("w:updateFields"))
    return el.get(qn("w:val")) if el is not None else None


class TestRendererStyle(unittest.TestCase):
    """TemplateStyle → DOCX 字段映射（用真实 general-thesis 样式）。"""

    def setUp(self):
        self.r = TemplateRenderer()
        self.tpl = load_basic("general-thesis")

    def test_fonts_cn_en(self):
        doc = self.r.render(self.tpl, make_spec())
        body_paras = [p for p in doc.paragraphs if p.text == "第一段正文内容。"]
        self.assertTrue(body_paras)
        run = body_paras[0].runs[0]
        self.assertEqual(run.font.name, "Times New Roman")
        self.assertEqual(_east_asia(run), "宋体")

    def test_font_size_bold_italic_underline(self):
        # heading1：黑体16 bold
        doc = self.r.render(self.tpl, make_spec())
        h1 = [p for p in doc.paragraphs if "引言" in p.text][0]
        run = h1.runs[0]
        self.assertAlmostEqual(run.font.size.pt, 16)
        self.assertTrue(run.font.bold)
        self.assertEqual(_east_asia(run), "黑体")
        # italic / underline 显式测试：构造样式
        style = TemplateStyle.from_dict({
            "font_size_pt": 14, "italic": True, "underline": True,
            "alignment": "center"})
        p = doc.add_paragraph()
        run = p.add_run("样式测试")
        self.r._apply_style(p, run, style)
        self.assertTrue(run.font.italic)
        self.assertTrue(run.font.underline)

    def test_alignment(self):
        doc = self.r.render(self.tpl, make_spec())
        title = [p for p in doc.paragraphs if p.text == "论文标题"][0]
        self.assertEqual(title.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        body = [p for p in doc.paragraphs if p.text == "第一段正文内容。"][0]
        self.assertEqual(body.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

    def test_line_spacing_multiple(self):
        doc = self.r.render(self.tpl, make_spec())
        body = [p for p in doc.paragraphs if p.text == "第一段正文内容。"][0]
        self.assertAlmostEqual(body.paragraph_format.line_spacing, 1.5)

    def test_line_spacing_exact(self):
        # references item 样式：exact 16
        doc = self.r.render(self.tpl, make_spec())
        ref = [p for p in doc.paragraphs if p.text.startswith("[1]")][0]
        pf = ref.paragraph_format
        self.assertEqual(pf.line_spacing_rule, WD_LINE_SPACING.EXACTLY)
        self.assertAlmostEqual(pf.line_spacing.pt, 16)

    def test_space_before_after(self):
        # title_zh: before 24 after 18
        doc = self.r.render(self.tpl, make_spec())
        title = [p for p in doc.paragraphs if p.text == "论文标题"][0]
        pf = title.paragraph_format
        self.assertAlmostEqual(pf.space_before.pt, 24)
        self.assertAlmostEqual(pf.space_after.pt, 18)

    def test_first_line_chars(self):
        # body 首行缩进 chars=2 → w:firstLineChars=200
        doc = self.r.render(self.tpl, make_spec())
        body = [p for p in doc.paragraphs if p.text == "第一段正文内容。"][0]
        self.assertEqual(_first_line_chars(body), "200")

    def test_first_line_pt(self):
        style = TemplateStyle.from_dict(
            {"first_line_indent": {"unit": "pt", "value": 24}})
        doc = self.r.render(make_template(), make_spec())
        p = doc.add_paragraph()
        run = p.add_run("x")
        self.r._apply_style(p, run, style)
        self.assertAlmostEqual(p.paragraph_format.first_line_indent.pt, 24)
        self.assertIsNone(_first_line_chars(p))

    def test_keep_with_next_page_break_before(self):
        style = TemplateStyle.from_dict(
            {"keep_with_next": True, "page_break_before": True})
        doc = self.r.render(make_template(), make_spec())
        p = doc.add_paragraph()
        run = p.add_run("x")
        self.r._apply_style(p, run, style)
        self.assertTrue(p.paragraph_format.keep_with_next)
        self.assertTrue(p.paragraph_format.page_break_before)


class TestRendererPage(unittest.TestCase):
    """页面尺寸 / 方向 / 页边距。"""

    def setUp(self):
        self.r = TemplateRenderer()

    def _page(self, **page):
        return make_template(page=page)

    def test_a4(self):
        doc = self.r.render(self._page(
            size="A4", margins={"top_mm": 25, "bottom_mm": 25,
                                "left_mm": 30, "right_mm": 25}), make_spec())
        sec = doc.sections[0]
        self.assertAlmostEqual(sec.page_width.mm, 210, places=1)
        self.assertAlmostEqual(sec.page_height.mm, 297, places=1)
        self.assertEqual(sec.orientation, WD_ORIENT.PORTRAIT)
        self.assertAlmostEqual(sec.top_margin.mm, 25, places=1)
        self.assertAlmostEqual(sec.left_margin.mm, 30, places=1)
        self.assertAlmostEqual(sec.bottom_margin.mm, 25, places=1)
        self.assertAlmostEqual(sec.right_margin.mm, 25, places=1)

    def test_custom_page_size(self):
        doc = self.r.render(self._page(size="210x280mm"), make_spec())
        sec = doc.sections[0]
        self.assertAlmostEqual(sec.page_width.mm, 210, places=1)
        self.assertAlmostEqual(sec.page_height.mm, 280, places=1)

    def test_landscape(self):
        doc = self.r.render(
            self._page(size="A4", orientation="landscape"), make_spec())
        sec = doc.sections[0]
        self.assertAlmostEqual(sec.page_width.mm, 297, places=1)
        self.assertAlmostEqual(sec.page_height.mm, 210, places=1)
        self.assertEqual(sec.orientation, WD_ORIENT.LANDSCAPE)

    def test_header_footer_distance(self):
        doc = self.r.render(self._page(
            header_distance_mm=15, footer_distance_mm=17.5), make_spec())
        sec = doc.sections[0]
        self.assertAlmostEqual(sec.header_distance.mm, 15, places=1)
        self.assertAlmostEqual(sec.footer_distance.mm, 17.5, places=1)

    def test_header_footer_content(self):
        tpl = make_template(
            header={"content": "论文页眉", "style": {
                "font_size_pt": 10.5, "alignment": "center"}},
            footer={"content": "- {page} -", "style": {
                "font_size_pt": 10.5, "alignment": "center"}})
        doc = self.r.render(tpl, make_spec())
        sec = doc.sections[0]
        header_text = "".join(r.text for r in sec.header.paragraphs[0].runs)
        self.assertIn("论文页眉", header_text)
        footer_p = sec.footer.paragraphs[0]
        footer_text = "".join(r.text for r in footer_p.runs)
        self.assertIn("- ", footer_text)
        self.assertIn(" -", footer_text)
        # PAGE 域存在
        instr = " ".join(_instr_texts(footer_p))
        self.assertIn("PAGE", instr)


class TestRendererContent(unittest.TestCase):
    """正文要素：标题/摘要/关键词/标题/段落/图表题注/参考文献。"""

    def setUp(self):
        self.r = TemplateRenderer()

    def test_title_abstract_keywords(self):
        doc = self.r.render(make_template(), make_spec())
        texts = [p.text for p in doc.paragraphs]
        self.assertIn("论文标题", texts)
        self.assertIn("摘  要", texts)
        self.assertIn("本文研究摘要内容。", texts)
        self.assertTrue(any(p.text.startswith("关键词：") for p in doc.paragraphs))

    def test_heading_levels_outline(self):
        tpl = load_basic("general-thesis")
        doc = self.r.render(tpl, make_spec())
        # h1 引言 / h2 研究背景 / h3 子问题
        h1 = [p for p in doc.paragraphs if "引言" in p.text][0]
        h2 = [p for p in doc.paragraphs if "研究背景" in p.text][0]
        h3 = [p for p in doc.paragraphs if "子问题" in p.text][0]
        self.assertEqual(_outline_level(h1), "0")
        self.assertEqual(_outline_level(h2), "1")
        self.assertEqual(_outline_level(h3), "2")
        # 一级标题样式（黑体16 bold）
        run = h1.runs[0]
        self.assertAlmostEqual(run.font.size.pt, 16)
        self.assertTrue(run.font.bold)
        self.assertEqual(_east_asia(run), "黑体")

    def test_numbering_enabled(self):
        doc = self.r.render(make_template(), make_spec())
        texts = [p.text for p in doc.paragraphs]
        self.assertIn("第一章 引言", texts)
        self.assertIn("1.1 研究背景", texts)
        self.assertIn("1.1.1 子问题", texts)
        self.assertIn("第二章 文献综述", texts)

    def test_numbering_disabled(self):
        numbering = {"enabled": False}
        doc = self.r.render(make_template(numbering=numbering), make_spec())
        texts = [p.text for p in doc.paragraphs]
        self.assertIn("引言", texts)
        self.assertNotIn("第一章 引言", texts)

    def test_numbering_no_double(self):
        # spec 已带编号 → 不重复编号
        spec = make_spec()
        spec["sections"][0]["text"] = "1 引言"
        doc = self.r.render(make_template(), spec)
        texts = [p.text for p in doc.paragraphs]
        self.assertIn("1 引言", texts)
        self.assertNotIn("第一章 1 引言", texts)

    def test_special_headings(self):
        # 致谢/参考文献 → 不自动编号，用对应 block
        spec = make_spec()
        spec["sections"] = [{"type": "h1", "text": "致谢"},
                            {"type": "p", "text": "感谢。"}]
        doc = self.r.render(make_template(), spec)
        ack = [p for p in doc.paragraphs if p.text == "致谢"][0]
        self.assertEqual(ack.text, "致谢")  # 无编号前缀
        self.assertEqual(_outline_level(ack), "0")

    def test_paragraph_body(self):
        doc = self.r.render(make_template(), make_spec())
        body = [p for p in doc.paragraphs if p.text == "第一段正文内容。"][0]
        run = body.runs[0]
        self.assertAlmostEqual(run.font.size.pt, 12)
        self.assertEqual(_east_asia(run), "宋体")
        self.assertEqual(_first_line_chars(body), "200")

    def test_pagebreak(self):
        spec = make_spec()
        spec["sections"] = [{"type": "p", "text": "第一页"},
                            {"type": "pagebreak"},
                            {"type": "p", "text": "第二页"}]
        doc = self.r.render(make_template(), spec)
        self.assertTrue(any(
            el.tag == qn("w:br") and el.get(qn("w:type")) == "page"
            for p in doc.paragraphs for el in p._p.iter()))

    def test_table_caption_and_table(self):
        spec = make_spec()
        spec["sections"] = [{
            "type": "table", "title": "表1 示例数据",
            "headers": ["变量", "均值"],
            "rows": [["A", "1.2"], ["B", "3.4"]]}]
        doc = self.r.render(make_template(), spec)
        # 题注
        cap = [p for p in doc.paragraphs if p.text == "表1 示例数据"][0]
        run = cap.runs[0]
        self.assertAlmostEqual(run.font.size.pt, 10.5)
        self.assertTrue(run.font.bold)
        self.assertEqual(cap.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        # 表格
        self.assertEqual(len(doc.tables), 1)
        self.assertEqual(len(doc.tables[0].rows), 3)

    def test_figure_and_caption(self):
        with tempfile.TemporaryDirectory() as tmp:
            img = Path(tmp) / "fig.png"
            img.write_bytes(_TINY_PNG)
            spec = make_spec()
            spec["sections"] = [{"type": "figure", "path": "fig.png",
                                 "title": "图1 示例"}]
            doc = self.r.render(make_template(), spec, base_dir=tmp)
            cap = [p for p in doc.paragraphs if p.text == "图1 示例"][0]
            run = cap.runs[0]
            self.assertAlmostEqual(run.font.size.pt, 10.5)
            self.assertTrue(run.font.bold)
            self.assertEqual(len(doc.inline_shapes), 1)

    def test_figure_missing_does_not_fail(self):
        spec = make_spec()
        spec["sections"] = [{"type": "figure", "path": "no-such.png",
                             "title": "图X"}]
        doc = self.r.render(make_template(), spec, base_dir=Path("."))
        self.assertTrue(doc.paragraphs)

    def test_references(self):
        # 用真实模板验证 item 样式（宋体 10.5 exact 16）
        tpl = load_basic("general-thesis")
        doc = self.r.render(tpl, make_spec())
        ref_title = [p for p in doc.paragraphs if p.text == "参考文献"]
        self.assertTrue(ref_title)
        ref = [p for p in doc.paragraphs if p.text.startswith("[1]")]
        self.assertTrue(ref)
        run = ref[0].runs[0]
        self.assertAlmostEqual(run.font.size.pt, 10.5)
        self.assertEqual(_east_asia(run), "宋体")
        pf = ref[0].paragraph_format
        self.assertEqual(pf.line_spacing_rule, WD_LINE_SPACING.EXACTLY)
        self.assertAlmostEqual(pf.line_spacing.pt, 16)

    def test_references_dict_entries_formatted(self):
        # dict 条目走现有 reference.py 格式化
        spec = make_spec()
        spec["sections"] = [{"type": "references", "items": [
            {"key": "a", "type": "journal", "title": "测试论文",
             "authors": ["张三"], "year": "2025", "source": "学报",
             "volume": "10", "issue": "2", "pages": "1-5"}]}]
        doc = self.r.render(make_template(), spec)
        texts = [p.text for p in doc.paragraphs]
        self.assertTrue(any("测试论文" in t for t in texts))
        self.assertTrue(any(t.startswith("[1]") for t in texts))


class TestRendererToc(unittest.TestCase):
    """TOC 域 + updateFields。"""

    def setUp(self):
        self.r = TemplateRenderer()

    def test_toc_field(self):
        doc = self.r.render(make_template(), make_spec())
        toc_paras = [p for p in doc.paragraphs
                     if any("TOC" in i for i in _instr_texts(p))]
        self.assertEqual(len(toc_paras), 1)
        self.assertEqual(_update_fields(doc), "true")

    def test_toc_disabled_block(self):
        # 用真实模板：禁用 toc 区块 → 不渲染 TOC
        tpl = load_basic("general-thesis")
        toc = tpl.get_block("toc")
        self.assertIsNotNone(toc)
        toc.enabled = False
        doc = self.r.render(tpl, make_spec())
        toc_paras = [p for p in doc.paragraphs
                     if any("TOC" in i for i in _instr_texts(p))]
        self.assertEqual(toc_paras, [])


class TestRendererFallback(unittest.TestCase):
    """缺失 block 兜底：当前模板 → 默认模板 → SAFE_DEFAULTS。"""

    def test_missing_blocks_safe_defaults(self):
        # 模板只有 title_zh 与 heading1
        data = make_template_dict()
        data["blocks"] = [data["blocks"][0]]  # 只留 title_zh
        tpl = Template.from_dict(data, template_id="min")
        doc = TemplateRenderer().render(tpl, make_spec())
        body = [p for p in doc.paragraphs if p.text == "第一段正文内容。"][0]
        # SAFE_DEFAULTS paragraph：宋体12 缩进2字符
        run = body.runs[0]
        self.assertAlmostEqual(run.font.size.pt, 12)
        self.assertEqual(_east_asia(run), "宋体")
        self.assertEqual(_first_line_chars(body), "200")
        # 标题区块也来自 SAFE_DEFAULTS
        title = [p for p in doc.paragraphs if p.text == "论文标题"][0]
        self.assertAlmostEqual(title.runs[0].font.size.pt, 22)

    def test_default_template_second_level(self):
        # 当前模板完全无 blocks → 用默认模板（general-thesis）的样式
        data = make_template_dict()
        data["blocks"] = []
        tpl = Template.from_dict(data, template_id="empty")
        default = load_basic("general-thesis")
        doc = TemplateRenderer(default_template=default).render(tpl, make_spec())
        body = [p for p in doc.paragraphs if p.text == "第一段正文内容。"][0]
        # general-thesis body 宋体 12
        self.assertEqual(_east_asia(body.runs[0]), "宋体")
        self.assertEqual(_first_line_chars(body), "200")


class TestRendererRealTemplates(unittest.TestCase):
    """三个真实 v2 模板渲染同一 spec，不抛异常且样式有差异。"""

    def setUp(self):
        self.r = TemplateRenderer()

    def test_all_basic_templates_render(self):
        spec = make_spec()
        for stem in ("general-thesis", "graduation-thesis", "course-paper"):
            tpl = load_basic(stem)
            doc = self.r.render(tpl, spec)
            self.assertTrue(doc.paragraphs, stem)
            self.assertTrue(any(p.text.startswith("第一章") for p in doc.paragraphs),
                            f"{stem} 编号失效")

    def test_template_block_differences(self):
        # 结构性差异：general 有 title_en 区块，course 无（英文标题默认关闭）
        general = load_basic("general-thesis")
        course = load_basic("course-paper")
        self.assertIsNotNone(general.get_block("title_en"))
        self.assertIsNone(course.get_block("title_en"))
        # 渲染英文标题：general 用自己的 title_en 样式（22pt）
        spec = make_spec()
        spec["meta"]["title_en"] = "Research Title"
        doc = self.r.render(general, spec)
        title_en = [p for p in doc.paragraphs if p.text == "Research Title"]
        self.assertTrue(title_en)
        self.assertAlmostEqual(title_en[0].runs[0].font.size.pt, 22)


class TestRendererDocument(unittest.TestCase):
    """落盘渲染。"""

    def test_render_document_saves(self):
        with tempfile.TemporaryDirectory() as tmp:
            out = TemplateRenderer().render_document(
                make_template(), make_spec(), tmp, "out.docx")
            self.assertTrue(out.is_file())
            self.assertGreater(out.stat().st_size, 0)
            from docx import Document as D
            doc = D(str(out))
            self.assertTrue(doc.paragraphs)


if __name__ == "__main__":
    unittest.main(verbosity=2)
