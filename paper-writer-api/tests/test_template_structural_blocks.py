from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from app.formatter import template_manager
from app.formatter.template import DEFAULT_TEMPLATES_ROOT, TemplateRenderer
from app.formatter.template.loader import TemplateLoader


class TestTemplateStructuralBlocks(unittest.TestCase):
    def test_extracts_table_and_section_metadata(self):
        with tempfile.TemporaryDirectory() as tmp:
            source_path = Path(tmp) / "template.docx"
            source = Document()
            table = source.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Method"
            table.cell(0, 1).text = "Purpose"
            table.cell(1, 0).text = "Review"
            table.cell(1, 1).text = "Framework"
            source.add_section(WD_SECTION.NEW_PAGE)
            source.save(source_path)
            config = template_manager._extract_config(source_path)
        self.assertEqual(len(config["tables"]), 1)
        self.assertEqual(config["tables"][0]["headers"], ["Method", "Purpose"])
        self.assertEqual(len(config["sections"]), 2)

    def test_renders_table_from_template_layout_and_section_boundary(self):
        template = TemplateLoader(DEFAULT_TEMPLATES_ROOT).load_template(
            DEFAULT_TEMPLATES_ROOT / "basic" / "general-thesis.json"
        )
        template.blocks.extend([
            type(template.blocks[0]).from_dict({
                "key": "table_1", "kind": "table", "enabled": True,
                "settings": {"headers": ["Method", "Purpose"], "column_widths_cm": [4.0, 8.0]},
            }),
            type(template.blocks[0]).from_dict({
                "key": "section_2", "kind": "sectionbreak", "enabled": True,
                "settings": {"start_type": "NEW_PAGE"},
            }),
        ])
        doc = TemplateRenderer().render(
            template,
            {
                "meta": {"toc": False},
                "sections": [
                    {"type": "table", "template_key": "table_1", "rows": [["Review", "Framework"]]},
                    {"type": "sectionbreak", "template_key": "section_2", "page_number": {"restart": True, "start": 1}},
                    {"type": "p", "text": "Body begins here."},
                ],
            },
        )
        self.assertEqual(len(doc.tables), 1)
        self.assertEqual([c.text for c in doc.tables[0].rows[0].cells], ["Method", "Purpose"])
        self.assertEqual(len(doc.sections), 2)
        pg_num = doc.sections[1]._sectPr.find(qn("w:pgNumType"))
        self.assertIsNotNone(pg_num)
        self.assertEqual(pg_num.get(qn("w:start")), "1")


if __name__ == "__main__":
    unittest.main()
