"""DocxExporter 文件输出层单元测试。

覆盖：Document → DOCX、自动创建目录、``.docx`` 后缀幂等、Windows 非法字符、
覆盖两态（overwrite True/False）、空文件名回退默认、超长文件名截断、
全非法字符回退默认、sanitize 幂等。

运行：``python -m unittest discover -s tests -v``
"""

from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document  # noqa: E402

from app.formatter.template.exporter import (  # noqa: E402
    DEFAULT_DOCX_FILENAME,
    DocxExporter,
)


def _blank_doc() -> Document:
    doc = Document()
    doc.add_paragraph("测试内容")
    return doc


class TestExportBasic(unittest.TestCase):
    """Document → DOCX 基础保存。"""

    def test_document_to_docx(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp) / "task"
            out = DocxExporter().export(_blank_doc(), task_dir)
            self.assertTrue(out.is_file())
            self.assertGreater(out.stat().st_size, 0)
            self.assertEqual(out.name, DEFAULT_DOCX_FILENAME)
            # 可用 python-docx 打开
            doc = Document(str(out))
            self.assertTrue(any(p.text == "测试内容" for p in doc.paragraphs))

    def test_returned_path_is_absolute(self):
        with tempfile.TemporaryDirectory() as tmp:
            out = DocxExporter().export(_blank_doc(), tmp)
            self.assertTrue(out.is_absolute())

    def test_default_filename_constant(self):
        self.assertEqual(DEFAULT_DOCX_FILENAME, "论文.docx")


class TestCreateDir(unittest.TestCase):
    """不存在 task_dir → 自动创建。"""

    def test_auto_create_nested_dir(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp) / "a" / "b" / "c"
            self.assertFalse(task_dir.exists())
            out = DocxExporter().export(_blank_doc(), task_dir)
            self.assertTrue(task_dir.is_dir())
            self.assertTrue(out.is_file())

    def test_existing_dir_reused(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp) / "exists"
            task_dir.mkdir()
            marker = task_dir / "keep.txt"
            marker.write_text("x", encoding="utf-8")
            DocxExporter().export(_blank_doc(), task_dir)
            self.assertTrue(marker.is_file())


class TestFilenameSuffix(unittest.TestCase):
    """.docx 后缀幂等。"""

    def test_no_suffix_gets_appended(self):
        self.assertEqual(DocxExporter().sanitize_filename("毕业论文"), "毕业论文.docx")

    def test_with_suffix_kept(self):
        self.assertEqual(DocxExporter().sanitize_filename("毕业论文.docx"), "毕业论文.docx")

    def test_uppercase_suffix_kept(self):
        self.assertEqual(DocxExporter().sanitize_filename("论文.DOCX"), "论文.DOCX")

    def test_sanitize_idempotent(self):
        ex = DocxExporter()
        for name in ("论文", "论文.docx", "我的 论文", "a:?b"):
            once = ex.sanitize_filename(name)
            self.assertEqual(ex.sanitize_filename(once), once)

    def test_export_without_suffix(self):
        with tempfile.TemporaryDirectory() as tmp:
            out = DocxExporter().export(_blank_doc(), tmp, "毕业论文")
            self.assertEqual(out.name, "毕业论文.docx")
            self.assertTrue(out.is_file())


class TestInvalidChars(unittest.TestCase):
    """Windows 非法字符处理。"""

    def test_windows_invalid_chars_replaced(self):
        ex = DocxExporter()
        # < > : " / \ | ? * 与控制字符
        name = ex.sanitize_filename('论文:测试?.docx')
        self.assertNotIn(":", name)
        self.assertNotIn("?", name)
        self.assertTrue(name.endswith(".docx"))
        self.assertNotIn("/", ex.sanitize_filename("a/b"))
        self.assertNotIn("\\", ex.sanitize_filename("a\\b"))
        self.assertNotIn("*", ex.sanitize_filename("a*b"))
        self.assertNotIn("|", ex.sanitize_filename("a|b"))
        self.assertNotIn('"', ex.sanitize_filename('a"b'))
        self.assertNotIn("<", ex.sanitize_filename("a<b"))
        self.assertNotIn(">", ex.sanitize_filename("a>b"))
        self.assertNotIn("\x00", ex.sanitize_filename("a\x00b"))

    def test_no_trailing_dot_or_space(self):
        ex = DocxExporter()
        self.assertFalse(ex.sanitize_filename("论文.").endswith("."))
        self.assertFalse(ex.sanitize_filename("论文 .").endswith(" ."))
        self.assertFalse(ex.sanitize_filename("论文 ").endswith(" "))

    def test_all_invalid_falls_back_default(self):
        ex = DocxExporter()
        # 全非法字符 → 全部替换为下划线，仍是安全非空文件名
        name = ex.sanitize_filename("::??**")
        self.assertNotIn(":", name)
        self.assertNotIn("?", name)
        self.assertTrue(name.endswith(".docx"))
        self.assertTrue(name.startswith("_"))
        # 结尾全为点 → 剥离后为空 → 默认文件名
        self.assertEqual(ex.sanitize_filename("..."), DEFAULT_DOCX_FILENAME)

    def test_export_safe_filename(self):
        with tempfile.TemporaryDirectory() as tmp:
            out = DocxExporter().export(_blank_doc(), tmp, "论文:测试?.docx")
            self.assertTrue(out.is_file())
            self.assertNotIn(":", out.name)
            self.assertNotIn("?", out.name)
            Document(str(out))


class TestEmptyFilename(unittest.TestCase):
    """空文件名 → 默认文件名（不生成 .docx）。"""

    def test_empty_falls_back_default(self):
        ex = DocxExporter()
        self.assertEqual(ex.sanitize_filename(""), DEFAULT_DOCX_FILENAME)
        self.assertEqual(ex.sanitize_filename("   "), DEFAULT_DOCX_FILENAME)
        self.assertEqual(ex.sanitize_filename(None), DEFAULT_DOCX_FILENAME)

    def test_export_empty_uses_default(self):
        with tempfile.TemporaryDirectory() as tmp:
            out = DocxExporter().export(_blank_doc(), tmp, "")
            self.assertEqual(out.name, DEFAULT_DOCX_FILENAME)
            self.assertNotEqual(out.name, ".docx")


class TestOverwrite(unittest.TestCase):
    """覆盖策略。"""

    def test_overwrite_true_replaces(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            out = DocxExporter().export(_blank_doc(), task_dir, "论文.docx")
            first_size = out.stat().st_size
            doc2 = Document()
            doc2.add_paragraph("第二版内容" * 50)
            out2 = DocxExporter().export(doc2, task_dir, "论文.docx", overwrite=True)
            self.assertEqual(out, out2)
            self.assertGreater(out2.stat().st_size, 0)
            # 内容已被替换
            reopened = Document(str(out2))
            self.assertTrue(any("第二版内容" in p.text for p in reopened.paragraphs))

    def test_overwrite_false_raises(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            DocxExporter().export(_blank_doc(), task_dir, "论文.docx")
            with self.assertRaises(FileExistsError):
                DocxExporter().export(_blank_doc(), task_dir, "论文.docx",
                                      overwrite=False)

    def test_overwrite_false_missing_ok(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            out = DocxExporter().export(_blank_doc(), task_dir, "新文档.docx",
                                        overwrite=False)
            self.assertTrue(out.is_file())

    def test_error_message_contains_path(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            DocxExporter().export(_blank_doc(), task_dir, "论文.docx")
            with self.assertRaises(FileExistsError) as ctx:
                DocxExporter().export(_blank_doc(), task_dir, "论文.docx",
                                      overwrite=False)
            self.assertIn("论文.docx", str(ctx.exception))


class TestLongFilename(unittest.TestCase):
    """超长文件名截断（保留 .docx 后缀）。"""

    def test_long_name_truncated(self):
        ex = DocxExporter()
        long_stem = "论" * 300
        name = ex.sanitize_filename(long_stem)
        self.assertLessEqual(len(name), 200)
        self.assertTrue(name.endswith(".docx"))
        self.assertEqual(len(name), 200)  # 满 200

    def test_long_with_suffix_truncated(self):
        ex = DocxExporter()
        name = ex.sanitize_filename("论" * 300 + ".docx")
        self.assertLessEqual(len(name), 200)
        self.assertTrue(name.endswith(".docx"))

    def test_short_name_untouched(self):
        ex = DocxExporter()
        name = ex.sanitize_filename("短文件名.docx")
        self.assertEqual(name, "短文件名.docx")


class TestRendererCompatibility(unittest.TestCase):
    """render_document 兼容入口仍走 Exporter。"""

    def test_render_document_uses_exporter(self):
        from unittest.mock import patch

        from app.formatter.template.renderer import TemplateRenderer
        from tests.test_template_renderer import make_spec, make_template

        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp) / "task"
            # renderer.py 在方法内局部导入，patch exporter 模块级名称即可拦截
            with patch(
                    "app.formatter.template.exporter.DocxExporter") as mock_cls:
                instance = mock_cls.return_value
                instance.export.return_value = task_dir / "out.docx"
                out = TemplateRenderer().render_document(
                    make_template(), make_spec(), task_dir, "out.docx")
                # 走的是 render() + exporter.export()
                self.assertEqual(out, task_dir / "out.docx")
                instance.export.assert_called_once()
            # patch 解除后真实保存仍可用
            out = TemplateRenderer().render_document(
                make_template(), make_spec(), task_dir, "out.docx")
            self.assertTrue(out.is_file())
            Document(str(out))


if __name__ == "__main__":
    unittest.main(verbosity=2)
