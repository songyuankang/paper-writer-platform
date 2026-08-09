"""TemplateRenderer 业务接入集成测试。

覆盖：现有论文数据（build_spec）→ spec 转换 → TemplateRenderer → 有效 DOCX
全链路，以及 format_paper 新链路（template_id）与旧链路（None）兼容、
模板选择（三个真实内置模板）、meta.toc 控制。

运行：``python -m unittest discover -s tests -v``
"""

from __future__ import annotations

import shutil
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

from app.config import settings  # noqa: E402
from app.db import init_db  # noqa: E402
from app.formatter.service import format_paper  # noqa: E402
from app.formatter.template import (  # noqa: E402
    DEFAULT_TEMPLATES_ROOT,
    TemplateRenderer,
    build_services,
    to_render_spec,
)
from app.formatter.template.loader import TemplateLoader  # noqa: E402
from app.models.generate import GenerateRequest  # noqa: E402
from app.services.content_generator import build_spec  # noqa: E402


def make_request(**overrides) -> GenerateRequest:
    req = GenerateRequest(
        title="集成测试论文", major="计算机科学与技术",
        paper_type="课程论文", word_count=3000,
        reference_style="gb7714")
    return req.model_copy(update=overrides) if overrides else req


def _instr_texts(p):
    return [el.text or "" for el in p._p.xpath(".//w:instrText")]


def _has_toc(doc) -> bool:
    return any("TOC" in i for p in doc.paragraphs for i in _instr_texts(p))


class IntegrationTestCase(unittest.TestCase):
    """临时 DB + 真实模板目录 + patch render_service.get_service。"""

    def setUp(self):
        self._tmp = tempfile.mkdtemp(prefix="renderint_")
        self._orig_db = settings.db_path
        settings.db_path = Path(self._tmp) / "test.db"
        init_db()
        _repo, _loader, service = build_services(DEFAULT_TEMPLATES_ROOT)
        self.service = service
        self._patcher = patch(
            "app.formatter.template.render_service.get_service",
            return_value=service)
        self._patcher.start()

    def tearDown(self):
        self._patcher.stop()
        settings.db_path = self._orig_db
        shutil.rmtree(self._tmp, ignore_errors=True)

    def _task_dir(self) -> Path:
        d = Path(self._tmp) / "task"
        d.mkdir(parents=True, exist_ok=True)
        return d


class TestSpecConverter(unittest.TestCase):
    """转换层：现有论文数据 → renderer spec。"""

    def test_build_spec_converts(self):
        # 真实数据生产者：build_spec(GenerateRequest)
        spec = build_spec(make_request())
        self.assertIn("sections", spec)
        self.assertIn("meta", spec)
        render_spec = to_render_spec(spec, paper_info={"title_en": "EN Title"})
        self.assertEqual(render_spec["meta"]["title"], "集成测试论文")
        self.assertEqual(render_spec["meta"]["title_en"], "EN Title")
        # sections 原样保留
        self.assertEqual(render_spec["sections"], spec["sections"])

    def test_unknown_section_preserved(self):
        spec = {"meta": {"title": "T"},
                "sections": [{"type": "future_type", "data": 1}]}
        out = to_render_spec(spec)
        self.assertEqual(out["sections"][0]["type"], "future_type")

    def test_toc_control(self):
        spec = build_spec(make_request())
        out = to_render_spec(spec, paper_info={"toc": False})
        self.assertFalse(out["meta"]["toc"])
        out2 = to_render_spec(spec, paper_info={})
        self.assertNotIn("toc", out2["meta"])

    def test_invalid_spec_skeleton(self):
        out = to_render_spec(None)
        self.assertEqual(out, {"meta": {}, "sections": [], "references": []})


class TestRenderChain(IntegrationTestCase):
    """现有论文数据 → renderer → 有效 DOCX。"""

    def render(self, template_id: str = "basic-general-thesis"):
        spec = build_spec(make_request())
        from app.formatter.template import render_service
        task_dir = self._task_dir()
        out = render_service.render_with_template(
            template_id, task_dir, spec, paper_info=None)
        return out, Document(str(out))

    def test_spec_render_docx_valid(self):
        # 真实数据（build_spec）标题自带编号 "1 引言" → 防重复编号，原样保留
        spec = build_spec(make_request())
        out, doc = self.render()
        self.assertTrue(out.is_file())
        self.assertGreater(out.stat().st_size, 0)
        texts = [p.text for p in doc.paragraphs]
        # 标题 / 摘要 / 关键词
        self.assertTrue(any("集成测试论文" in t for t in texts))
        self.assertTrue(any("摘  要" in t for t in texts))
        self.assertTrue(any(t.startswith("关键词：") for t in texts))
        # 已带编号的 h1 保持原文（不重复编号）
        self.assertTrue(any(t == "1 引言" or t.startswith("1 引言") for t in texts))
        self.assertFalse(any(t.startswith("第一章 1 引言") for t in texts))
        # 正文
        self.assertTrue(any(p.text and "本文" in p.text for p in doc.paragraphs))
        # 参考文献
        self.assertTrue(any("参考文献" in t for t in texts))
        # TOC
        self.assertTrue(_has_toc(doc))

    def test_unnumbered_input_gets_numbered(self):
        # 未编号标题 → renderer 按模板 numbering 生成 第一章 / 1.1
        spec = {
            "meta": {"title": "编号测试"},
            "sections": [
                {"type": "h1", "text": "引言"},
                {"type": "p", "text": "正文。"},
                {"type": "h2", "text": "背景"},
            ],
        }
        task_dir = self._task_dir()
        from app.formatter.template import render_service
        out = render_service.render_with_template(
            "basic-general-thesis", task_dir, spec)
        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs]
        self.assertIn("第一章 引言", texts)
        self.assertIn("1.1 背景", texts)

    def test_template_selection_three_basic(self):
        spec = build_spec(make_request())
        for tid in ("basic-general-thesis", "basic-graduation-thesis",
                    "basic-course-paper"):
            out, doc = self.render(tid)
            self.assertTrue(out.is_file(), tid)
            self.assertTrue(doc.paragraphs, tid)
            texts = [p.text for p in doc.paragraphs]
            self.assertTrue(any(t == "1 引言" or t.startswith("1 引言")
                                for t in texts), tid)

    def test_renderer_meta_toc_false(self):
        spec = build_spec(make_request())
        spec["meta"]["toc"] = False
        loader = TemplateLoader(DEFAULT_TEMPLATES_ROOT)
        tpl = loader.load_template(
            DEFAULT_TEMPLATES_ROOT / "basic" / "general-thesis.json")
        doc = TemplateRenderer().render(tpl, spec)
        self.assertFalse(_has_toc(doc))


class TestFormatPaperIntegration(IntegrationTestCase):
    """format_paper 新/旧链路。"""

    def test_new_path_template_id(self):
        task_dir = self._task_dir()
        spec = build_spec(make_request())
        files = format_paper("t1", task_dir, make_request().model_dump(),
                             spec, template_id="basic-general-thesis")
        # docx 生成
        self.assertTrue((task_dir / "论文.docx").is_file())
        self.assertTrue((task_dir / "paper_spec.json").is_file())
        self.assertTrue((task_dir / "references.json").is_file())
        self.assertIn("论文.docx", [f.split("/")[-1] for f in files])
        doc = Document(str(task_dir / "论文.docx"))
        texts = [p.text for p in doc.paragraphs]
        # 自带编号标题保持原文（防重复）
        self.assertTrue(any(t == "1 引言" or t.startswith("1 引言")
                            for t in texts))
        self.assertTrue(_has_toc(doc))
        self.assertTrue(any("集成测试论文" in t for t in texts))

    def test_new_path_graduation(self):
        task_dir = self._task_dir()
        spec = build_spec(make_request())
        format_paper("t2", task_dir, make_request().model_dump(), spec,
                     template_id="basic-graduation-thesis")
        self.assertTrue((task_dir / "论文.docx").is_file())
        doc = Document(str(task_dir / "论文.docx"))
        texts = [p.text for p in doc.paragraphs]
        self.assertTrue(any(t == "1 引言" or t.startswith("1 引言")
                            for t in texts))

    def test_old_path_no_template_id(self):
        task_dir = self._task_dir()
        spec = build_spec(make_request())
        files = format_paper("t3", task_dir, make_request().model_dump(), spec)
        self.assertTrue((task_dir / "论文.docx").is_file())
        self.assertIn("论文.docx", [f.split("/")[-1] for f in files])
        # 旧链路仍能出 docx（向后兼容）
        Document(str(task_dir / "论文.docx"))

    def test_new_path_unknown_template_falls_back_default(self):
        # 模板 id 不存在 → resolve 回退默认模板，仍出 docx
        task_dir = self._task_dir()
        spec = build_spec(make_request())
        format_paper("t4", task_dir, make_request().model_dump(), spec,
                     template_id="basic-not-exist")
        self.assertTrue((task_dir / "论文.docx").is_file())


if __name__ == "__main__":
    unittest.main(verbosity=2)
