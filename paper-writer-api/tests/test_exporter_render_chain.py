"""DocxExporter / RenderService 完整链路测试。

覆盖：
1. ``render_with_template`` 完整链路：TemplateService.resolve → to_render_spec
   → TemplateRenderer.render（Document）→ DocxExporter.export（.docx）→ Path，
   且最终 DOCX 可用 ``Document(path)`` 打开；
2. 三个真实内置模板（basic-general-thesis / basic-graduation-thesis /
   basic-course-paper）分别生成 DOCX：文件存在 / 大小 > 0 / 可打开 /
   标题存在 / 正文存在；
3. 至少一个完整真实链路：template_id → TemplateService → RenderService →
   TemplateRenderer → DocxExporter → 论文.docx。

运行：``python -m unittest discover -s tests -v``
"""

from __future__ import annotations

import shutil
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document  # noqa: E402
from docx.document import Document as DocumentClass  # noqa: E402

from app.config import settings  # noqa: E402
from app.db import init_db  # noqa: E402
from app.formatter.template import (  # noqa: E402
    DEFAULT_TEMPLATES_ROOT,
    DEFAULT_DOCX_FILENAME,
    DocxExporter,
    TemplateRenderer,
    build_services,
    to_render_spec,
)
from app.formatter.template.loader import TemplateLoader  # noqa: E402
from app.models.generate import GenerateRequest  # noqa: E402
from app.services.content_generator import build_spec  # noqa: E402

REAL_TEMPLATES = (
    "basic-general-thesis",
    "basic-graduation-thesis",
    "basic-course-paper",
)


def make_request(**overrides) -> GenerateRequest:
    req = GenerateRequest(
        title="输出层链路测试论文", major="计算机科学与技术",
        paper_type="课程论文", word_count=3000,
        reference_style="gb7714")
    return req.model_copy(update=overrides) if overrides else req


def make_spec() -> dict:
    """未编号标题的 spec（触发 renderer 编号），验证标题/正文。"""
    return {
        "meta": {
            "title": "输出层链路测试论文",
            "abstract": "本文用于验证 DocxExporter 与 RenderService 完整链路。",
            "keywords": ["测试", "链路"],
            "reference_style": "gb7714",
        },
        "sections": [
            {"type": "h1", "text": "引言"},
            {"type": "p", "text": "第一段正文内容，用于验证正文存在。"},
            {"type": "h2", "text": "研究背景"},
            {"type": "p", "text": "第二段正文内容。"},
            {"type": "references", "items": [
                "[1] 张三. 示例研究[J]. 期刊, 2025, 10(2): 1-5."]},
        ],
        "references": [
            {"id": 1, "authors": "张三", "title": "示例研究",
             "journal": "期刊", "year": "2025",
             "volume": "10", "issue": "2", "pages": "1-5",
             "type": "journal"}],
    }


class RenderServiceChainTestCase(unittest.TestCase):
    """临时 DB + 真实模板目录 + patch render_service.get_service。"""

    def setUp(self):
        self._tmp = tempfile.mkdtemp(prefix="exporter_")
        self._orig_db = settings.db_path
        settings.db_path = Path(self._tmp) / "test.db"
        init_db()
        _repo, _loader, service = build_services(DEFAULT_TEMPLATES_ROOT)
        self.service = service
        self._patcher = patch(
            "app.formatter.template.render_service.get_service",
            return_value=service)
        self._patcher.start()

    def tearDown(self):
        self._patcher.stop()
        settings.db_path = self._orig_db
        shutil.rmtree(self._tmp, ignore_errors=True)

    def _task_dir(self, name: str) -> Path:
        d = Path(self._tmp) / "tasks" / name
        d.mkdir(parents=True, exist_ok=True)
        return d

    def test_render_service_full_chain(self):
        """render_with_template 完整链路 → 有效 DOCX。"""
        from app.formatter.template import render_service
        task_dir = self._task_dir("full_chain")
        spec = make_spec()
        out = render_service.render_with_template(
            "basic-general-thesis", task_dir, spec, paper_info=None)
        # Path 返回、默认文件名、存在、>0
        self.assertIsInstance(out, Path)
        self.assertEqual(out.name, DEFAULT_DOCX_FILENAME)
        self.assertTrue(out.is_file())
        self.assertGreater(out.stat().st_size, 0)
        # Document(path) 可打开，标题/正文存在
        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs]
        self.assertTrue(any("输出层链路测试论文" in t for t in texts))
        self.assertTrue(any("正文" in t for t in texts))

    def test_render_service_chain_respects_out_name(self):
        from app.formatter.template import render_service
        task_dir = self._task_dir("out_name")
        out = render_service.render_with_template(
            "basic-general-thesis", task_dir, make_spec(),
            out_name="毕业论文")
        self.assertEqual(out.name, "毕业论文.docx")
        self.assertTrue(out.is_file())

    def test_renderer_render_returns_document_no_save(self):
        """TemplateRenderer.render 只返回内存 Document，不落盘。"""
        loader = TemplateLoader(DEFAULT_TEMPLATES_ROOT)
        tpl = loader.load_template(
            DEFAULT_TEMPLATES_ROOT / "basic" / "general-thesis.json")
        spec = to_render_spec(make_spec())
        doc = TemplateRenderer().render(tpl, spec)
        self.assertIsInstance(doc, DocumentClass)
        self.assertTrue(doc.paragraphs)
        # 目录中不应出现任何 .docx（render 不落盘）
        task_dir = self._task_dir("no_save")
        docs = list(task_dir.rglob("*.docx"))
        self.assertEqual(docs, [])


class RealTemplateChainTestCase(RenderServiceChainTestCase):
    """真实模板真实链路：template_id → Service → RenderService → Renderer → Exporter → 论文.docx。"""

    def _assert_docx_valid(self, out: Path, tid: str) -> None:
        self.assertTrue(out.is_file(), f"{tid}: 文件不存在")
        self.assertGreater(out.stat().st_size, 0, f"{tid}: 文件为空")
        doc = Document(str(out))  # python-docx 可打开
        texts = [p.text for p in doc.paragraphs]
        self.assertTrue(any("输出层链路测试论文" in t for t in texts),
                        f"{tid}: 标题缺失")
        self.assertTrue(any("正文" in t for t in texts),
                        f"{tid}: 正文缺失")
        self.assertTrue(any("参考文献" in t for t in texts),
                        f"{tid}: 参考文献缺失")

    def test_three_real_templates_generate_docx(self):
        for tid in REAL_TEMPLATES:
            with self.subTest(template=tid):
                task_dir = self._task_dir(f"real_{tid}")
                from app.formatter.template import render_service
                out = render_service.render_with_template(
                    tid, task_dir, make_spec(), paper_info=None)
                self._assert_docx_valid(out, tid)
                self.assertEqual(out.name, DEFAULT_DOCX_FILENAME)

    def test_full_real_chain_to_docx_file(self):
        """完整真实链路：template_id → TemplateService → RenderService
        → TemplateRenderer → DocxExporter → 论文.docx。"""
        # 1. TemplateService.resolve（真实 service + 真实模板目录）
        template = self.service.resolve("basic-graduation-thesis")
        self.assertIsNotNone(template)
        # 2. to_render_spec
        render_spec = to_render_spec(make_spec(), paper_info=None)
        # 3. TemplateRenderer.render → Document（内存）
        renderer = TemplateRenderer(default_template=self.service.default_template())
        doc = renderer.render(template, render_spec)
        self.assertIsInstance(doc, DocumentClass)
        # 4. DocxExporter.export → .docx 文件（任务目录此时尚未创建，验证自动建目录）
        task_dir = Path(self._tmp) / "real_full_chain"
        shutil.rmtree(task_dir, ignore_errors=True)  # 确保目录不存在
        self.assertFalse(task_dir.exists())
        out = DocxExporter().export(doc, task_dir, DEFAULT_DOCX_FILENAME)
        # 5. 校验
        self.assertEqual(out, task_dir / DEFAULT_DOCX_FILENAME)
        self.assertTrue(task_dir.is_dir())  # Exporter 自动建目录
        self._assert_docx_valid(out, "basic-graduation-thesis")

    def test_export_overwrite_false_on_real_output(self):
        """真实链路输出后，overwrite=False 应抛 FileExistsError。"""
        task_dir = self._task_dir("overwrite")
        from app.formatter.template import render_service
        render_service.render_with_template(
            "basic-course-paper", task_dir, make_spec())
        target = task_dir / DEFAULT_DOCX_FILENAME
        self.assertTrue(target.is_file())
        with self.assertRaises(FileExistsError):
            DocxExporter().export(Document(), task_dir,
                                  DEFAULT_DOCX_FILENAME, overwrite=False)


if __name__ == "__main__":
    unittest.main(verbosity=2)
