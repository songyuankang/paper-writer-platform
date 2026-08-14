"""Regression tests for semantic heading style preservation."""

from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from app.formatter.template import DEFAULT_TEMPLATES_ROOT, TemplateRenderer
from app.formatter.template.loader import TemplateLoader


class TestSemanticHeadingStyles(unittest.TestCase):
    def test_renderer_reuses_heading_styles_observed_in_source_template(self):
        template = TemplateLoader(DEFAULT_TEMPLATES_ROOT).load_template(
            DEFAULT_TEMPLATES_ROOT / "basic" / "general-thesis.json"
        )
        with tempfile.TemporaryDirectory() as tmp:
            source_path = Path(tmp) / "source.docx"
            source = Document()
            source.add_paragraph("Source chapter", style="Heading 3")
            source.add_paragraph("Source section", style="Heading 4")
            source.add_paragraph("Source subsection", style="Heading 5")
            source.save(source_path)

            doc = TemplateRenderer().render(
                template,
                {
                    "meta": {"toc": False},
                    "sections": [
                        {"type": "h1", "text": "Chapter"},
                        {"type": "h2", "text": "Section"},
                        {"type": "h3", "text": "Subsection"},
                    ],
                },
                source_docx=source_path,
            )

        chapter = next(p for p in doc.paragraphs if p.text.endswith("Chapter"))
        section = next(p for p in doc.paragraphs if p.text.endswith("Section"))
        subsection = next(p for p in doc.paragraphs if p.text.endswith("Subsection"))
        self.assertEqual(chapter.style.name, "Heading 3")
        self.assertEqual(section.style.name, "Heading 4")
        self.assertEqual(subsection.style.name, "Heading 5")


    def test_renderer_normalizes_legacy_hyphenated_heading_numbers(self):
        template = TemplateLoader(DEFAULT_TEMPLATES_ROOT).load_template(
            DEFAULT_TEMPLATES_ROOT / "basic" / "general-thesis.json"
        )
        doc = TemplateRenderer().render(
            template,
            {
                "meta": {"toc": False},
                "sections": [
                    {"type": "h1", "text": "Chapter"},
                    {"type": "h2", "text": "Section"},
                    {"type": "h3", "text": "1-1.1 Legacy title"},
                ],
            },
        )
        legacy = next(p.text for p in doc.paragraphs if p.text.endswith("Legacy title"))
        self.assertEqual(legacy, "1.1.1 Legacy title")
        self.assertNotIn("0.0.", legacy)
    def test_renderer_creates_semantic_fallback_when_source_style_is_latent(self):
        template = TemplateLoader(DEFAULT_TEMPLATES_ROOT).load_template(
            DEFAULT_TEMPLATES_ROOT / "basic" / "general-thesis.json"
        )
        doc = TemplateRenderer().render(
            template,
            {
                "meta": {"toc": False},
                "sections": [{"type": "h1", "text": "Chapter"}],
            },
        )
        paragraph = next(p for p in doc.paragraphs if p.text.endswith("Chapter"))
        self.assertNotEqual(paragraph.style.name, "Normal")
        self.assertTrue(paragraph.style.name.startswith(("Heading", "PW Heading")))
if __name__ == "__main__":
    unittest.main()