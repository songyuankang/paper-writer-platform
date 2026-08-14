from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.formatter.template import DEFAULT_TEMPLATES_ROOT, TemplateRenderer
from app.formatter.template.loader import TemplateLoader
from app.formatter.template import render_service


def load_template():
    return TemplateLoader(DEFAULT_TEMPLATES_ROOT).load_template(
        DEFAULT_TEMPLATES_ROOT / "basic" / "general-thesis.json"
    )


class TestTemplatePrelimPagination(unittest.TestCase):
    def test_abstract_titles_center_and_prelim_pages_separate(self):
        doc = TemplateRenderer().render(load_template(), {
            "meta": {"abstract": "Chinese abstract.", "keywords": ["one"],
                     "abstract_en": "English abstract.", "keywords_en": ["one"],
                     "toc": True},
            "sections": [{"type": "h1", "text": "Chapter 1"}],
        })
        chinese = next(p for p in doc.paragraphs if p.text == "摘  要")
        english = next(p for p in doc.paragraphs if p.text == "Abstract")
        self.assertEqual(chinese.alignment, 1)
        self.assertEqual(english.alignment, 1)
        self.assertGreaterEqual(sum('w:type="page"' in p._p.xml for p in doc.paragraphs), 3)

    def test_footer_replaces_inherited_page_field(self):
        with tempfile.TemporaryDirectory() as tmp:
            source_path = Path(tmp) / "source.docx"
            source = Document()
            footer = source.sections[0].footer.paragraphs[0]
            footer.add_run("-")
            run = footer.add_run()
            begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
            instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
            end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
            run._r.extend([begin, instr, end])
            footer.add_run("-")
            source.save(source_path)
            template = load_template()
            template.footer = {"content": "-{page}-", "style": {"alignment": "center"}}
            doc = TemplateRenderer().render(template, {"meta": {"toc": False}, "sections": []}, source_docx=source_path)
        self.assertEqual(doc.sections[0].footer.paragraphs[0]._p.xml.count(" PAGE "), 1)

    @patch("app.formatter.template.render_service.deepseek.chat")
    def test_missing_english_abstract_is_generated(self, mocked_chat):
        mocked_chat.return_value = "Abstract: An English abstract.\nKeywords: tomato, organic farming"
        spec = {"meta": {"abstract": "中文摘要", "keywords": ["番茄"]}, "sections": []}
        render_service._ensure_english_abstract(spec)
        self.assertEqual(spec["meta"]["abstract_en"], "An English abstract.")
        self.assertEqual(spec["meta"]["keywords_en"], ["tomato", "organic farming"])

    def test_english_abstract_and_keywords_render_once(self):
        doc = TemplateRenderer().render(load_template(), {
            "meta": {
                "abstract": "Chinese abstract.",
                "keywords": ["one"],
                "abstract_en": "English abstract.",
                "keywords_en": ["one"],
            },
            "sections": [],
        })
        texts = [paragraph.text.strip() for paragraph in doc.paragraphs]
        self.assertEqual(sum(text == "Abstract" for text in texts), 1)
        self.assertEqual(sum(text.startswith("Keywords:") for text in texts), 1)


if __name__ == "__main__":
    unittest.main()
